from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime
from typing import List, Optional
import asyncpg
import gc
import random
from fastapi import Request, HTTPException

def create_title_style(document):
    """Create and return a title style for the document."""
    style = document.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    return style

def create_heading_style(document, level, size):
    """Create and return a heading style for the document."""
    style_name = f'CustomHeading{level}'
    style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(size)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    return style

def create_normal_style(document):
    """Create and return a normal text style for the document."""
    style = document.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    return style

def create_code_style(doc):
    """Create a custom style for code blocks."""
    style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(9)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.left_indent = Pt(18)
    paragraph_format.right_indent = Pt(18)
    paragraph_format.background_color = RGBColor(240, 240, 240)

def add_title_page(document, title, author, date):
    """Add a title page to the document."""
    # Add a page break to ensure title page is separate
    document.add_page_break()
    
    # Add title
    title_para = document.add_paragraph(title)
    title_para.style = 'CustomTitle'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    document.add_paragraph()
    document.add_paragraph()
    
    # Add author
    author_para = document.add_paragraph(f"Author: {author}")
    author_para.style = 'CustomNormal'
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = document.add_paragraph(date)
    date_para.style = 'CustomNormal'
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break after title page
    document.add_page_break()

def add_table_of_contents(document):
    """Add a table of contents to the document."""
    heading = document.add_paragraph("Table of Contents")
    heading.style = 'CustomHeading1'
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph()
    
    # Add table of contents entries
    toc_entries = [
        ("1. Introduction", 1),
        ("2. Problem Statement", 1),
        ("3. Project Requirements", 1),
        ("4. Approach and Methodology", 1),
        ("5. Technical Implementation", 1),
        ("6. Code Explanation", 1),
        ("7. Results and Examples", 1),
        ("8. Future Work", 1),
        ("9. Conclusion", 1),
        ("10. References", 1),
    ]
    
    for entry, level in toc_entries:
        para = document.add_paragraph(entry)
        para.style = 'CustomNormal'
        para.paragraph_format.left_indent = Inches(0.5 * level)
    
    document.add_page_break()

def add_section(document, title, content, level=1):
    """Add a section to the document."""
    heading = document.add_paragraph(title)
    heading.style = f'CustomHeading{level}'
    
    for paragraph in content:
        para = document.add_paragraph(paragraph)
        para.style = 'CustomNormal'

def add_code_snippet(document, code, language="python"):
    """Add a code snippet to the document."""
    # Clean up the code to ensure XML compatibility
    code = code.strip()
    code = code.replace('\u0000', '')  # Remove NULL bytes
    code = ''.join(char for char in code if ord(char) >= 32 or char in '\n\r\t')  # Remove control characters
    
    para = document.add_paragraph()
    para.style = 'Code'
    run = para.add_run(f"{code}")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

def add_hyperlink(document, text, url):
    """Add a hyperlink to the document."""
    para = document.add_paragraph()
    para.style = 'CustomNormal'
    run = para.add_run(f"{text}: {url}")
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
    run.font.underline = True
    run.hyperlink = url

def add_code_documentation(document):
    """Add detailed code documentation section to the report."""
    document.add_heading("6.4 Code Documentation", level=2)
    
    # Document translation_bot.py
    document.add_heading("translation_bot.py", level=3)
    document.add_paragraph(
        "The translation_bot.py file serves as the core of our application, implementing the FastAPI server "
        "and handling translation requests. This file contains the main application logic, API endpoints, "
        "and integration with various AI models."
    )

    # Main FastAPI Application
    add_code_subsection(
        document,
        "FastAPI Application Setup",
        """The application is initialized using FastAPI, setting up CORS middleware and route handlers:
        - Creates FastAPI instance with custom title and description
        - Configures CORS middleware to handle cross-origin requests
        - Sets up WebSocket connection for real-time translation updates
        - Implements rate limiting and request validation
        - Handles API key management and security

        Key implementation details:
        - Uses environment variables for configuration
        - Implements custom middleware for request logging
        - Handles API key rotation and validation
        - Manages concurrent request processing
        - Provides detailed error responses"""
    )

    # Add code example for FastAPI setup
    add_code_snippet(
        document,
        """app = FastAPI(
    title="PersianAI Translation Service",
    description="Professional Persian to English translation service",
    version="1.0.0"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add custom middleware
@app.middleware("http")
async def add_process_time_header(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    response.headers["X-Process-Time"] = str(process_time)
    return response""",
        "python"
    )

    # Translation Logic
    add_code_subsection(
        document,
        "Translation Core Logic",
        """The translation logic includes:
        - Integration with Azure Translator API for accurate translations
        - Custom text preprocessing for Persian language
        - Handling of special characters and formatting
        - Caching mechanism for improved performance
        - Support for multiple AI models (OpenAI and Gemini)
        - Text chunking for handling large documents

        Implementation features:
        - Smart text chunking based on sentence boundaries
        - Context preservation between chunks
        - Parallel processing of chunks
        - Result aggregation and formatting
        - Quality validation of translations"""
    )

    # Add code example for translation logic
    add_code_snippet(
        document,
        """async def translate_text(text: str, model_type: str = "gpt-3.5-turbo") -> str:
    \"\"\"Translate Persian text to English using specified model.\"\"\"
    try:
        # Preprocess text
        text = preprocess_persian_text(text)
        
        # Split into chunks if needed
        chunks = await chunk_text(text)
        
        # Process chunks in parallel
        tasks = []
        for chunk in chunks:
            if model_type in ["gpt-3.5-turbo", "gpt-4"]:
                task = process_openai_translation(chunk, model_type)
            else:
                task = process_gemini_translation(chunk)
            tasks.append(task)
        
        # Wait for all translations
        translations = await asyncio.gather(*tasks)
        
        # Combine and validate results
        result = combine_translations(translations)
        validate_translation(result)
        
        return result
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        raise""",
        "python"
    )

    # Document changes.py
    document.add_heading("changes.py", level=3)
    document.add_paragraph(
        "The changes.py file manages version control and tracks modifications to translations. "
        "It implements a sophisticated diff system that can detect and visualize changes at the word level, "
        "making it easy to review and validate translations."
    )

    add_code_subsection(
        document,
        "Change Tracking System",
        """Features of the change tracking system:
        - Version history maintenance
        - Diff generation between versions
        - Rollback capability for translations
        - Audit trail for all modifications
        - Word-level change detection
        - HTML-based diff visualization
        - Change statistics and metrics

        Key components:
        - Sequence matching algorithm for word-level diffs
        - HTML generation with color coding
        - Change statistics calculation
        - Version comparison utilities
        - Rollback functionality"""
    )

    # Add code example for change tracking
    add_code_snippet(
        document,
        """def track_changes(original: str, modified: str) -> Dict[str, Any]:
    \"\"\"Track changes between original and modified text.\"\"\"
    # Split into words while preserving formatting
    orig_words = split_into_words(original)
    mod_words = split_into_words(modified)
    
    # Create sequence matcher
    matcher = SequenceMatcher(None, orig_words, mod_words)
    
    # Process changes
    changes = {
        'replacements': 0,
        'deletions': 0,
        'insertions': 0,
        'details': []
    }
    
    # Analyze each operation
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'replace':
            changes['replacements'] += 1
            changes['details'].append({
                'type': 'replace',
                'original': ' '.join(orig_words[i1:i2]),
                'modified': ' '.join(mod_words[j1:j2])
            })
        elif op == 'delete':
            changes['deletions'] += 1
            changes['details'].append({
                'type': 'delete',
                'original': ' '.join(orig_words[i1:i2])
            })
        elif op == 'insert':
            changes['insertions'] += 1
            changes['details'].append({
                'type': 'insert',
                'modified': ' '.join(mod_words[j1:j2])
            })
    
    return changes""",
        "python"
    )

    # Document improvements.py
    document.add_heading("improvements.py", level=3)
    document.add_paragraph(
        "The improvements.py file contains logic for enhancing translation quality through various "
        "optimization techniques and custom prompts. It implements sophisticated text processing "
        "algorithms to improve the accuracy and naturalness of translations."
    )

    add_code_subsection(
        document,
        "Translation Enhancement Features",
        """Key improvement features include:
        - Context-aware translation refinement
        - Persian language-specific optimizations
        - Quality metrics calculation
        - Suggestion system for better translations
        - Custom prompt templates for different content types
        - Error correction and grammar improvement

        Implementation details:
        - Custom prompt engineering for different content types
        - Context window management
        - Quality scoring system
        - Grammar correction rules
        - Style consistency checks"""
    )

    # Add code example for improvements
    add_code_snippet(
        document,
        """def enhance_translation(text: str, context: Dict[str, Any] = None) -> str:
    \"\"\"Enhance translation quality using context and rules.\"\"\"
    try:
        # Apply Persian-specific rules
        text = apply_persian_rules(text)
        
        # Generate context-aware prompt
        prompt = generate_context_prompt(text, context)
        
        # Get model suggestions
        suggestions = get_model_suggestions(prompt)
        
        # Apply improvements
        improved_text = apply_improvements(text, suggestions)
        
        # Validate quality
        quality_score = calculate_quality_score(improved_text)
        
        if quality_score < QUALITY_THRESHOLD:
            # Apply additional improvements if needed
            improved_text = apply_additional_improvements(improved_text)
        
        return improved_text
    except Exception as e:
        logger.error(f"Enhancement error: {str(e)}")
        raise""",
        "python"
    )

    # Document test files
    document.add_heading("Test Files", level=3)
    document.add_paragraph(
        "The project includes a comprehensive test suite spread across multiple files, "
        "each focusing on different aspects of the system. These tests ensure reliability "
        "and maintainability of the codebase."
    )

    # Document test_changes.py
    document.add_heading("test_changes.py", level=4)
    document.add_paragraph(
        "The test_changes.py file contains unit tests for the change tracking system, "
        "ensuring accurate detection and visualization of text modifications."
    )

    add_code_subsection(
        document,
        "Testing Framework",
        """Testing features include:
        - Unit tests for text difference detection
        - Tests for HTML diff generation
        - Validation of change statistics
        - Edge case handling tests
        - Performance benchmarks
        - Error handling verification

        Test coverage:
        - Basic text comparison
        - Complex formatting scenarios
        - Unicode character handling
        - Large text processing
        - Error conditions"""
    )

    # Document test_connections.py
    document.add_heading("test_connections.py", level=4)
    document.add_paragraph(
        "The test_connections.py file verifies API connectivity and integration, "
        "ensuring reliable communication with external services."
    )

    add_code_subsection(
        document,
        "API Testing Features",
        """API testing capabilities include:
        - Connection validation for all external APIs
        - Response format verification
        - Error handling testing
        - Rate limit testing
        - Timeout handling verification
        - API key validation tests

        Test scenarios:
        - Successful API calls
        - Error responses
        - Rate limiting
        - Network issues
        - Authentication failures"""
    )

    # Document test_gemini.py
    document.add_heading("test_gemini.py", level=4)
    document.add_paragraph(
        "The test_gemini.py file specifically tests Google's Gemini model integration, "
        "ensuring proper functionality of the Gemini-based translation features."
    )

    add_code_subsection(
        document,
        "Gemini Model Testing",
        """Testing features for Gemini integration:
        - Model initialization verification
        - Response quality assessment
        - Configuration validation
        - Error handling for model-specific issues
        - Performance monitoring
        - Resource cleanup verification

        Test aspects:
        - Model loading
        - Response formatting
        - Error scenarios
        - Resource management
        - Performance metrics"""
    )

def add_code_subsection(document, title, content):
    """Helper function to add a code documentation subsection."""
    document.add_heading(title, level=4)
    document.add_paragraph(content)
    document.add_paragraph()  # Add spacing between sections

def add_function_documentation(doc):
    """Add detailed documentation for each function in the project."""
    doc.add_heading('6.5 Function Documentation', level=1)
    
    # Translation Bot Functions
    doc.add_heading('6.5.1 Translation Bot Functions', level=2)
    
    # Core Functions
    doc.add_heading('Core Functions', level=3)
    doc.add_paragraph('The core functions handle the main application setup and configuration:')
    
    doc.add_paragraph('create_app() Function:')
    doc.add_paragraph('Purpose: Creates and configures the FastAPI application with all necessary middleware and settings.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Initializes FastAPI with custom title and description\n• Sets up CORS middleware for cross-origin requests\n• Configures rate limiting and request validation\n• Establishes WebSocket connection handling\n• Sets up API key management and security')
    add_code_snippet(doc, '''
    def create_app():
        """Creates and configures the FastAPI application.
        
        Returns:
            FastAPI: Configured FastAPI application instance.
        """
    ''')
    
    doc.add_paragraph('add_process_time_header() Function:')
    doc.add_paragraph('Purpose: Adds processing time information to response headers for performance monitoring.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Measures request processing time\n• Adds timing information to response headers\n• Helps with performance monitoring\n• Enables response time tracking')
    add_code_snippet(doc, '''
    def add_process_time_header(request: Request, call_next):
        """Adds processing time header to response.
        
        Args:
            request (Request): FastAPI request object
            call_next: Next middleware function
            
        Returns:
            Response: Response with processing time header
        """
    ''')
    
    # Translation Functions
    doc.add_heading('Translation Functions', level=3)
    doc.add_paragraph('The translation functions handle the core translation logic:')
    
    doc.add_paragraph('translate_text() Function:')
    doc.add_paragraph('Purpose: Main translation endpoint that processes translation requests.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Handles text translation requests\n• Supports multiple AI models\n• Manages text chunking for large inputs\n• Implements caching for performance\n• Provides real-time progress updates')
    add_code_snippet(doc, '''
    async def translate_text(request: Request):
        """Handles text translation requests.
        
        Args:
            request (Request): FastAPI request object containing text and model type
            
        Returns:
            dict: Contains translated text
        """
    ''')
    
    doc.add_paragraph('process_openai_translation() Function:')
    doc.add_paragraph('Purpose: Processes translations using OpenAI models.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Integrates with OpenAI API\n• Handles model-specific parameters\n• Manages API rate limiting\n• Implements error handling\n• Provides fallback options')
    add_code_snippet(doc, '''
    async def process_openai_translation(text: str, model: str) -> str:
        """Processes translation using OpenAI API.
        
        Args:
            text (str): Text to translate
            model (str): OpenAI model to use
            
        Returns:
            str: Translated text
        """
    ''')
    
    doc.add_paragraph('process_gemini_translation() Function:')
    doc.add_paragraph('Purpose: Processes translations using Google\'s Gemini model.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Integrates with Gemini API\n• Handles model-specific configurations\n• Manages API responses\n• Implements error handling\n• Provides model-specific optimizations')
    add_code_snippet(doc, '''
    async def process_gemini_translation(text: str, model: str) -> str:
        """Processes translation using Gemini API.
        
        Args:
            text (str): Text to translate
            model (str): Gemini model to use
            
        Returns:
            str: Translated text
        """
    ''')
    
    # Text Processing Functions
    doc.add_heading('Text Processing Functions', level=3)
    doc.add_paragraph('The text processing functions handle text manipulation and optimization:')
    
    doc.add_paragraph('chunk_text() Function:')
    doc.add_paragraph('Purpose: Splits large text into manageable chunks for processing.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Smart text splitting based on sentence boundaries\n• Maintains context between chunks\n• Handles special characters\n• Preserves formatting\n• Optimizes chunk size')
    add_code_snippet(doc, '''
    async def chunk_text(text: str, chunk_size: int) -> List[str]:
        """Splits text into chunks of specified size.
        
        Args:
            text (str): Text to split
            chunk_size (int): Maximum size of each chunk
            
        Returns:
            List[str]: List of text chunks
        """
    ''')
    
    # Changes Module Functions
    doc.add_heading('6.5.2 Changes Module Functions', level=2)
    doc.add_paragraph('The changes module functions handle version control and change tracking:')
    
    doc.add_paragraph('process_gemini_edit() Function:')
    doc.add_paragraph('Purpose: Processes text editing using Gemini model.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Handles text editing requests\n• Integrates with Gemini API\n• Tracks changes made\n• Generates change statistics\n• Provides edit history')
    add_code_snippet(doc, '''
    async def process_gemini_edit(text: str) -> str:
        """Processes text editing using Gemini API.
        
        Args:
            text (str): Text to edit
            
        Returns:
            str: Edited text
        """
    ''')
    
    doc.add_paragraph('process_openai_edit() Function:')
    doc.add_paragraph('Purpose: Processes text editing using OpenAI models.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Handles text editing requests\n• Integrates with OpenAI API\n• Tracks modifications\n• Generates change reports\n• Maintains edit history')
    add_code_snippet(doc, '''
    async def process_openai_edit(text: str, model: str = ModelType.GPT35.value) -> str:
        """Processes text editing using OpenAI API.
        
        Args:
            text (str): Text to edit
            model (str): OpenAI model to use
            
        Returns:
            str: Edited text
        """
    ''')
    
    # Improvements Module
    doc.add_heading('6.5.3 Improvements Module', level=2)
    doc.add_paragraph('The improvements module contains prompt templates and enhancement logic:')
    
    doc.add_paragraph('TRANSLATION_PROMPT:')
    doc.add_paragraph('Purpose: Template for Persian to English translation.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Optimized for Persian language\n• Handles cultural context\n• Maintains text formatting\n• Preserves meaning\n• Supports multiple content types')
    
    doc.add_paragraph('EDIT_PROMPT:')
    doc.add_paragraph('Purpose: Template for basic Persian text editing.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Basic text improvements\n• Grammar corrections\n• Style enhancements\n• Format preservation\n• Context awareness')
    
    doc.add_paragraph('EDIT_PROMPT_DETAILED:')
    doc.add_paragraph('Purpose: Template for comprehensive Persian text editing.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Detailed text analysis\n• Advanced improvements\n• Style optimization\n• Format enhancement\n• Context preservation')
    
    # Test Functions
    doc.add_heading('6.5.4 Test Functions', level=2)
    doc.add_paragraph('The test functions ensure system reliability and functionality:')
    
    # Test Changes Functions
    doc.add_heading('Test Changes Functions', level=3)
    doc.add_paragraph('Functions for testing the change tracking system:')
    
    doc.add_paragraph('generate_test_text() Function:')
    doc.add_paragraph('Purpose: Generates sample Persian text for testing.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Creates test data\n• Simulates real text\n• Tests various scenarios\n• Validates functionality\n• Measures performance')
    add_code_snippet(doc, '''
    async def generate_test_text(word_count: int) -> str:
        """Generates sample Persian text for testing.
        
        Args:
            word_count (int): Number of words to generate
            
        Returns:
            str: Generated test text
        """
    ''')
    
    doc.add_paragraph('test_chunking() Function:')
    doc.add_paragraph('Purpose: Tests the text chunking mechanism.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Validates chunk size\n• Tests boundary conditions\n• Checks context preservation\n• Verifies formatting\n• Measures performance')
    add_code_snippet(doc, '''
    async def test_chunking() -> None:
        """Tests the text chunking mechanism with various sizes.
        
        Tests different word counts and chunk sizes to ensure proper text splitting.
        """
    ''')
    
    # Test Connections Functions
    doc.add_heading('Test Connections Functions', level=3)
    doc.add_paragraph('Functions for testing API connectivity:')
    
    doc.add_paragraph('test_openai_connection() Function:')
    doc.add_paragraph('Purpose: Tests connection to OpenAI API.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Validates API keys\n• Tests connectivity\n• Checks response format\n• Verifies error handling\n• Measures response time')
    add_code_snippet(doc, '''
    async def test_openai_connection() -> bool:
        """Tests connection to OpenAI API.
        
        Returns:
            bool: True if connection successful, False otherwise
        """
    ''')
    
    doc.add_paragraph('test_gemini_connection() Function:')
    doc.add_paragraph('Purpose: Tests connection to Gemini API.')
    doc.add_paragraph('Key Features:')
    doc.add_paragraph('• Validates API keys\n• Tests connectivity\n• Checks response format\n• Verifies error handling\n• Measures response time')
    add_code_snippet(doc, '''
    async def test_gemini_connection() -> bool:
        """Tests connection to Gemini API.
        
        Returns:
            bool: True if connection successful, False otherwise
        """
    ''')

    # API Client Management
    doc.add_heading('API Client Management Functions', level=3)
    
    doc.add_paragraph('get_sync_openai_client() Function:')
    doc.add_paragraph('Purpose: Initializes and caches a synchronous OpenAI client with optimized settings.')
    add_code_snippet(doc, '''
@lru_cache()
def get_sync_openai_client():
    """Get or create synchronous OpenAI client instance."""
    try:
        return OpenAI(
            api_key=OPENAI_API_KEY,
            timeout=httpx.Timeout(60.0, connect=10.0, read=30.0, write=30.0),
            max_retries=3
        )
    except Exception as e:
        logger.error(f"Error initializing sync OpenAI client: {str(e)}")
        raise
''')

    doc.add_paragraph('get_async_openai_client() Function:')
    doc.add_paragraph('Purpose: Initializes and caches an asynchronous OpenAI client for concurrent operations.')
    add_code_snippet(doc, '''
@lru_cache()
def get_async_openai_client():
    """Get or create async OpenAI client instance."""
    try:
        return AsyncOpenAI(
            api_key=OPENAI_API_KEY,
            timeout=httpx.Timeout(60.0, connect=10.0, read=30.0, write=30.0),
            max_retries=3
        )
    except Exception as e:
        logger.error(f"Error initializing async OpenAI client: {str(e)}")
        raise
''')

    doc.add_paragraph('validate_api_keys() Function:')
    doc.add_paragraph('Purpose: Validates the presence of required API keys with descriptive error messages.')
    add_code_snippet(doc, '''
def validate_api_keys():
    """Validate API keys and raise descriptive errors."""
    missing_keys = []
    if not OPENAI_API_KEY:
        missing_keys.append("OpenAI API key")
    if not GEMINI_API_KEY:
        missing_keys.append("Gemini API key")
    
    if missing_keys:
        raise ValueError(f"Missing required API keys: {', '.join(missing_keys)}")
''')

    doc.add_paragraph('initialize_api_clients() Function:')
    doc.add_paragraph('Purpose: Initializes and tests API clients to ensure they are working properly.')
    add_code_snippet(doc, '''
async def initialize_api_clients():
    """Initialize API clients and verify they work."""
    try:
        # Test OpenAI client
        client = get_sync_openai_client()
        models = client.models.list()
        logger.info("OpenAI client initialized successfully")
        logger.info(f"Available OpenAI models: {[model.id for model in models.data]}")

        # Test Gemini model
        model = get_gemini_model()
        logger.info("Gemini model initialized successfully")
        
        return True
    except Exception as e:
        logger.error(f"Failed to initialize API clients: {str(e)}")
        return False
''')

    # Model Types and Configuration
    doc.add_heading('Model Types and Configuration', level=3)
    
    doc.add_paragraph('ModelType Class:')
    doc.add_paragraph('Purpose: Defines available AI models and their characteristics.')
    add_code_snippet(doc, '''
class ModelType(str, Enum):
    GPT35 = "gpt-3.5-turbo"
    GPT4 = "gpt-4"
    GEMINI_FLASH = "models/gemini-1.5-flash-8b"
    GEMINI_PRO = "models/gemini-1.5-pro-latest"
    
    @property
    def description(self) -> str:
        descriptions = {
            self.GPT35: "Fast and reliable processing with good accuracy",
            self.GPT4: "Most accurate processing, better understanding of context and nuances",
            self.GEMINI_FLASH: "Gemini 1.5 Flash - Fast and efficient model for text processing",
            self.GEMINI_PRO: "Gemini 1.5 Pro - Advanced model with better understanding"
        }
        return descriptions.get(self, "Unknown model")
        
    @property
    def max_tokens(self) -> int:
        limits = {
            self.GPT35: 30000,
            self.GPT4: 50000,
            self.GEMINI_FLASH: 1_000_000,
            self.GEMINI_PRO: 2000000
        }
        return limits.get(self, 30000)
''')

    doc.add_paragraph('EditMode Class:')
    doc.add_paragraph('Purpose: Defines text editing modes (fast vs. detailed).')
    add_code_snippet(doc, '''
class EditMode(str, Enum):
    FAST = "fast"
    DETAILED = "detailed"

class EditRequest(BaseModel):
    text: str
    model: str
    mode: EditMode = EditMode.FAST

class EditResponse(BaseModel):
    edited_text: str
    technical_explanation: str
''')

    # Text Processing Functions
    doc.add_heading('Additional Text Processing Functions', level=3)
    
    doc.add_paragraph('validate_word_count() Function:')
    doc.add_paragraph('Purpose: Ensures edited text length stays within acceptable bounds.')
    add_code_snippet(doc, '''
def validate_word_count(original_text: str, edited_text: str, tolerance: int = 50) -> bool:
    """Validate that edited text word count is within tolerance of original."""
    original_count = len(original_text.split())
    edited_count = len(edited_text.split())
    return abs(original_count - edited_count) <= tolerance
''')

    doc.add_paragraph('process_gemini_edit() Function:')
    doc.add_paragraph('Purpose: Processes text editing requests using the Gemini model.')
    add_code_snippet(doc, '''
async def process_gemini_edit(text: str, mode: EditMode = EditMode.FAST) -> str:
    """Process text editing using Gemini model."""
    try:
        model = get_gemini_model()
        prompt = EDIT_PROMPT if mode == EditMode.FAST else EDIT_PROMPT_DETAILED
        response = await model.generate_content(f"{prompt}\n\nText: {text}")
        return response.text
    except Exception as e:
        logger.error(f"Gemini editing error: {str(e)}")
        raise
''')

    doc.add_paragraph('process_openai_edit() Function:')
    doc.add_paragraph('Purpose: Processes text editing requests using OpenAI models.')
    doc.add_paragraph('Purpose: Processes text editing requests using OpenAI models.')
    add_code_snippet(doc, '''
async def process_openai_edit(text: str, model: str = ModelType.GPT35.value, 
                            mode: EditMode = EditMode.FAST) -> str:
    """Process text editing using OpenAI model."""
    try:
        client = get_async_openai_client()
        prompt = EDIT_PROMPT if mode == EditMode.FAST else EDIT_PROMPT_DETAILED
        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"OpenAI editing error: {str(e)}")
        raise
''')

    doc.add_paragraph('is_title() Function:')
    doc.add_paragraph('Purpose: Detects if a text segment is likely a title based on characteristics.')
    add_code_snippet(doc, '''
def is_title(text: str, is_paragraph_start: bool = False) -> bool:
    """Determine if text is likely a title."""
    # Title characteristics
    characteristics = [
        len(text.split()) <= 10,  # Short length
        text.strip().endswith((':', '؛', '.')),  # Ends with certain punctuation
        text.isupper(),  # All caps (for English)
        any(char.isdigit() for char in text[:2]),  # Starts with number
        is_paragraph_start  # At start of paragraph
    ]
    return any(characteristics)
''')

def add_project_configuration(doc):
    """Add documentation about project configuration and environment setup."""
    doc.add_heading('6.6 Project Configuration', level=1)
    
    # Environment Configuration
    doc.add_heading('Environment Configuration', level=2)
    doc.add_paragraph('The project uses environment variables for configuration management:')
    add_code_snippet(doc, '''
# .env.template
OPENAI_API_KEY=your_openai_api_key
GEMINI_API_KEY=your_gemini_api_key
MODEL_TYPE=gpt-3.5-turbo
DEBUG_MODE=False
LOG_LEVEL=INFO
''')
    
    # Dependencies
    doc.add_heading('Dependencies', level=2)
    doc.add_paragraph('Key project dependencies and their purposes:')
    add_code_snippet(doc, '''
# requirements.txt
fastapi==0.68.1
uvicorn==0.15.0
python-dotenv==0.19.0
openai==1.0.0
google-generativeai==0.3.0
python-multipart==0.0.5
aiohttp==3.8.1
''')
    
    # Deployment Configuration
    doc.add_heading('Deployment Configuration', level=2)
    doc.add_paragraph('The project includes deployment configurations for various platforms:')
    add_code_snippet(doc, '''
# Procfile
web: uvicorn translation_bot:app --host=0.0.0.0 --port=${PORT:-8000}
''')

def add_testing_framework(doc):
    """Add detailed documentation about the testing framework."""
    doc.add_heading('6.7 Testing Framework', level=1)
    
    # Test Organization
    doc.add_heading('Test Organization', level=2)
    doc.add_paragraph('The project includes three main test files:')
    
    # Test Changes
    doc.add_heading('test_changes.py', level=3)
    doc.add_paragraph('Tests for the change tracking system:')
    add_code_snippet(doc, '''
async def test_chunking() -> None:
    """Tests the text chunking mechanism."""
    test_text = "This is a test sentence. " * 10
    chunks = await chunk_text(test_text)
    assert len(chunks) > 0
    assert all(len(chunk) <= MAX_CHUNK_SIZE for chunk in chunks)
''')
    
    # Test Connections
    doc.add_heading('test_connections.py', level=3)
    doc.add_paragraph('Tests for API connectivity:')
    add_code_snippet(doc, '''
async def test_openai_connection() -> bool:
    """Tests OpenAI API connectivity."""
    try:
        client = get_sync_openai_client()
        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": "Test"}]
        )
        return True
    except Exception as e:
        logger.error(f"OpenAI connection test failed: {str(e)}")
        return False
''')
    
    # Test Gemini
    doc.add_heading('test_gemini.py', level=3)
    doc.add_paragraph('Tests for Gemini model integration:')
    add_code_snippet(doc, '''
def test_gemini_connection() -> bool:
    """Tests Gemini API connectivity."""
    try:
        model = get_gemini_model()
        response = model.generate_content("Test")
        return True
    except Exception as e:
        logger.error(f"Gemini connection test failed: {str(e)}")
        return False
''')

def add_static_assets(doc):
    """Add documentation about static assets and templates."""
    doc.add_heading('6.8 Static Assets and Templates', level=1)
    
    # Templates
    doc.add_heading('HTML Templates', level=2)
    doc.add_paragraph('The project includes HTML templates for the web interface:')
    add_code_snippet(doc, '''
<!-- templates/index.html -->
<!DOCTYPE html>
<html>
<head>
    <title>PersianAI Translation</title>
    <link rel="stylesheet" href="/static/css/style.css">
</head>
<body>
    <div class="container">
        <h1>PersianAI Translation</h1>
        <div id="translation-form">
            <!-- Translation form content -->
        </div>
    </div>
</body>
</html>
''')
    
    # Static Files
    doc.add_heading('Static Files', level=2)
    doc.add_paragraph('The project includes CSS and JavaScript files for styling and interactivity:')
    add_code_snippet(doc, '''
/* static/css/style.css */
.container {
    max-width: 1200px;
    margin: 0 auto;
    padding: 20px;
}

.translation-form {
    background: #f5f5f5;
    padding: 20px;
    border-radius: 5px;
}
''')

def add_logging_monitoring(doc):
    """Add documentation about logging and monitoring."""
    doc.add_heading('6.9 Logging and Monitoring', level=1)
    
    # Logging Configuration
    doc.add_heading('Logging Configuration', level=2)
    doc.add_paragraph('The project implements comprehensive logging:')
    add_code_snippet(doc, '''
# Logging configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('server.log'),
        logging.StreamHandler()
    ]
)
''')
    
    # Performance Monitoring
    doc.add_heading('Performance Monitoring', level=2)
    doc.add_paragraph('The project includes performance monitoring:')
    add_code_snippet(doc, '''
@app.middleware("http")
async def add_process_time_header(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    response.headers["X-Process-Time"] = str(process_time)
    return response
''')

def add_all_python_files(doc):
    """Add documentation for all Python files in the project."""
    doc.add_heading('6.10 Complete Python Files Overview', level=1)
    
    # Core Application Files
    doc.add_heading('Core Application Files', level=2)
    
    # translation_bot.py
    doc.add_heading('translation_bot.py', level=3)
    doc.add_paragraph('Main application file containing:')
    add_code_snippet(doc, '''
# Key components:
- FastAPI server implementation
- API endpoints for translation and editing
- OpenAI and Gemini model integration
- WebSocket communication
- Rate limiting and error handling
- Text chunking and processing
- API client management
- Middleware configuration
''')
    
    # changes.py
    doc.add_heading('changes.py', level=3)
    doc.add_paragraph('Change tracking and version control:')
    add_code_snippet(doc, '''
# Key features:
- Word-level change detection
- HTML diff generation
- Change statistics calculation
- Version history management
- Text comparison utilities
- Change visualization
''')
    
    # improvements.py
    doc.add_heading('improvements.py', level=3)
    doc.add_paragraph('Translation enhancement features:')
    add_code_snippet(doc, '''
# Main components:
- Translation prompt templates
- Editing prompt templates
- Context-aware refinement
- Quality metrics calculation
- Persian language rules
- Style consistency checks
''')
    
    # Test Files
    doc.add_heading('Test Files', level=2)
    
    # test_changes.py
    doc.add_heading('test_changes.py', level=3)
    doc.add_paragraph('Unit tests for change tracking:')
    add_code_snippet(doc, '''
# Test coverage:
- Text difference detection
- HTML generation
- Change statistics
- Word count validation
- Edge case handling
- Performance benchmarks
''')
    
    # test_connections.py
    doc.add_heading('test_connections.py', level=3)
    doc.add_paragraph('API connectivity tests:')
    add_code_snippet(doc, '''
# Test features:
- OpenAI API connection
- Gemini API connection
- Rate limit testing
- Error handling
- Response validation
- Timeout handling
''')
    
    # test_gemini.py
    doc.add_heading('test_gemini.py', level=3)
    doc.add_paragraph('Gemini model integration tests:')
    add_code_snippet(doc, '''
# Test aspects:
- Model initialization
- Response generation
- Error scenarios
- Configuration validation
- Performance testing
- Resource cleanup
''')
    
    # Configuration Files
    doc.add_heading('Configuration Files', level=2)
    
    # .env.template and .env.example
    doc.add_heading('Environment Configuration', level=3)
    doc.add_paragraph('Environment variable templates:')
    add_code_snippet(doc, '''
# Key configurations:
- API keys management
- Model selection
- Debug settings
- Log levels
- Performance parameters
- Security settings
''')
    
    # requirements.txt
    doc.add_heading('Dependencies', level=3)
    doc.add_paragraph('Project dependencies:')
    add_code_snippet(doc, '''
# Main dependencies:
- fastapi==0.68.1
- uvicorn==0.15.0
- python-dotenv==0.19.0
- openai==1.0.0
- google-generativeai==0.3.0
- python-multipart==0.0.5
- aiohttp==3.8.1
''')
    
    # Procfile
    doc.add_heading('Deployment Configuration', level=3)
    doc.add_paragraph('Deployment settings:')
    add_code_snippet(doc, '''
# Deployment config:
web: uvicorn translation_bot:app --host=0.0.0.0 --port=${PORT:-8000}
''')

def add_detailed_code_explanations(doc):
    """Add detailed code explanations for each Python file."""
    doc.add_heading('6.11 Detailed Code Explanations', level=1)
    
    # translation_bot.py
    doc.add_heading('translation_bot.py', level=2)
    doc.add_paragraph('Core application file with detailed implementation:')
    add_code_snippet(doc, '''
# Key Components and Their Implementation:

1. FastAPI Application Setup:
   - Creates FastAPI instance with custom title and description
   - Configures CORS middleware for cross-origin requests
   - Sets up WebSocket connection for real-time updates
   - Implements rate limiting and request validation

2. API Client Management:
   - OpenAI client initialization with caching
   - Gemini model setup with optimized configuration
   - Error handling and retry mechanisms
   - API key validation and rotation

3. Translation Core Logic:
   - Text preprocessing for Persian language
   - Smart chunking based on sentence boundaries
   - Parallel processing of text chunks
   - Result aggregation and validation

4. Error Handling:
   - Comprehensive exception handling
   - Detailed error logging
   - Graceful degradation
   - User-friendly error messages

5. Performance Optimization:
   - Request time tracking
   - Response caching
   - Resource cleanup
   - Memory management
''')
    
    # changes.py
    doc.add_heading('changes.py', level=2)
    doc.add_paragraph('Change tracking system implementation:')
    add_code_snippet(doc, '''
# Implementation Details:

1. Word-Level Change Detection:
   - Splits text into words while preserving formatting
   - Uses SequenceMatcher for efficient comparison
   - Handles Persian-specific characters
   - Maintains word order and context

2. HTML Diff Generation:
   - Creates color-coded HTML output
   - Highlights additions, deletions, and changes
   - Preserves text formatting and structure
   - Supports RTL text direction

3. Change Statistics:
   - Tracks number of changes by type
   - Calculates change percentages
   - Maintains version history
   - Generates change summaries

4. Text Processing:
   - Handles special characters
   - Preserves formatting
   - Manages text direction
   - Supports multiple languages
''')
    
    # improvements.py
    doc.add_heading('improvements.py', level=2)
    doc.add_paragraph('Translation enhancement implementation:')
    add_code_snippet(doc, '''
# Key Features:

1. Prompt Templates:
   - TRANSLATION_PROMPT: Optimized for Persian to English translation
   - EDIT_PROMPT: Basic text editing template
   - EDIT_PROMPT_DETAILED: Comprehensive editing template

2. Context-Aware Processing:
   - Maintains context between chunks
   - Preserves document structure
   - Handles formatting and style
   - Manages terminology consistency

3. Quality Enhancement:
   - Grammar correction
   - Style improvement
   - Terminology standardization
   - Readability optimization

4. Persian Language Rules:
   - Grammar rules implementation
   - Style guidelines
   - Formatting standards
   - Cultural considerations
''')
    
    # Test Files
    doc.add_heading('Test Files Implementation', level=2)
    
    # test_changes.py
    doc.add_heading('test_changes.py', level=3)
    doc.add_paragraph('Change tracking test implementation:')
    add_code_snippet(doc, '''
# Test Coverage:

1. Text Processing Tests:
   - Word splitting functionality
   - Character encoding handling
   - Format preservation
   - Direction handling

2. Change Detection Tests:
   - Word-level comparison
   - HTML generation
   - Statistics calculation
   - Edge case handling

3. Performance Tests:
   - Large text processing
   - Memory usage
   - Processing speed
   - Resource cleanup
''')
    
    # test_connections.py
    doc.add_heading('test_connections.py', level=3)
    doc.add_paragraph('API connectivity test implementation:')
    add_code_snippet(doc, '''
# Test Implementation:

1. API Connection Tests:
   - OpenAI API connectivity
   - Gemini API connectivity
   - Authentication validation
   - Response verification

2. Error Handling Tests:
   - Rate limit handling
   - Timeout scenarios
   - Invalid responses
   - Network issues

3. Performance Tests:
   - Response times
   - Connection stability
   - Resource usage
   - Error recovery
''')
    
    # test_gemini.py
    doc.add_heading('test_gemini.py', level=3)
    doc.add_paragraph('Gemini model test implementation:')
    add_code_snippet(doc, '''
# Test Features:

1. Model Integration Tests:
   - Model initialization
   - Configuration validation
   - Response generation
   - Error handling

2. Performance Tests:
   - Response times
   - Resource usage
   - Memory management
   - Connection stability

3. Quality Tests:
   - Output validation
   - Format preservation
   - Error recovery
   - Resource cleanup
''')

def add_api_endpoints_documentation(doc):
    """Add documentation for all API endpoints."""
    doc.add_heading('6.12 API Endpoints Documentation', level=1)
    
    # Root Endpoint
    doc.add_heading('Root Endpoint', level=2)
    doc.add_paragraph('Main translation endpoint:')
    add_code_snippet(doc, '''
    # POST /
    - Purpose: Main translation endpoint
    - Request Body:
      - text: Text to translate
      - model_type: "openai" or "gemini"
      - edit_mode: "basic" or "detailed"
    - Response:
      - translated_text: Translated content
      - changes: Change statistics
      - html_diff: Visual diff of changes
    ''')
    
    # Edit Endpoint
    doc.add_heading('Edit Endpoint', level=2)
    doc.add_paragraph('Text editing endpoint:')
    add_code_snippet(doc, '''
    # POST /edit
    - Purpose: Edit and improve text
    - Request Body:
      - text: Text to edit
      - model_type: "openai" or "gemini"
      - edit_mode: "basic" or "detailed"
    - Response:
      - edited_text: Improved text
      - changes: Change statistics
      - html_diff: Visual diff of changes
    ''')
    
    # WebSocket Endpoint
    doc.add_heading('WebSocket Endpoint', level=2)
    doc.add_paragraph('Real-time updates endpoint:')
    add_code_snippet(doc, '''
    # WebSocket /ws
    - Purpose: Real-time translation updates
    - Events:
      - translation_progress: Progress updates
      - translation_complete: Final results
      - error: Error notifications
    - Data:
      - progress: Completion percentage
      - current_chunk: Current processing chunk
      - status: Operation status
    ''')
    
    # Health Check Endpoint
    doc.add_heading('Health Check Endpoint', level=2)
    doc.add_paragraph('System health monitoring:')
    add_code_snippet(doc, '''
    # GET /health
    - Purpose: System health check
    - Response:
      - status: "healthy" or "unhealthy"
      - api_status: API connectivity status
      - model_status: Model availability
      - system_metrics: Performance metrics
    ''')

def add_error_handling_documentation(doc):
    """Add error handling and troubleshooting documentation section"""
    doc.add_heading('6.13 Error Handling and Troubleshooting', level=1)
    
    # Common Error Types
    doc.add_heading('Common Error Types', level=2)
    doc.add_paragraph('The application handles various types of errors:')
    
    # API Errors
    doc.add_heading('API Errors', level=3)
    doc.add_paragraph('• OpenAI API errors (rate limiting, authentication)\n• Gemini API errors\n• Network connectivity issues')
    
    # Processing Errors
    doc.add_heading('Processing Errors', level=3)
    doc.add_paragraph('• Translation errors\n• Memory overflow errors\n• Input validation errors')
    
    # System Errors
    doc.add_heading('System Errors', level=3)
    doc.add_paragraph('• Configuration errors\n• Database errors\n• Resource exhaustion')

def add_performance_optimization_documentation(doc):
    """Add performance optimization documentation section"""
    doc.add_heading('6.14 Performance Optimization', level=1)
    
    # Text Processing Optimization
    doc.add_heading('Text Processing Optimization', level=2)
    doc.add_paragraph('The application implements several optimization strategies for text processing:')
    add_code_snippet(doc, '''
    # Text chunking for large inputs
    async def chunk_text(text: str, max_chunk_size: int = 1000) -> List[str]:
        """Split text into manageable chunks while preserving context."""
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in text.split('.'):
            sentence_size = len(sentence)
            if current_size + sentence_size <= max_chunk_size:
                current_chunk.append(sentence)
                current_size += sentence_size
            else:
                chunks.append('.'.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_size
                
        if current_chunk:
            chunks.append('.'.join(current_chunk))
            
        return chunks
    ''')
    
    # Memory Management
    doc.add_heading('Memory Management', level=2)
    doc.add_paragraph('Memory optimization strategies include:')
    add_code_snippet(doc, '''
    # Resource cleanup
    def cleanup_resources():
        gc.collect()  # Garbage collection
        
    # Memory-efficient processing
    async def process_large_text(text: str):
        chunks = await chunk_text(text)
        results = []
        
        for chunk in chunks:
            result = await process_chunk(chunk)
            results.append(result)
            await cleanup_resources()  # Clean up after each chunk
            
        return combine_results(results)
    ''')
    
    # API Request Optimization
    doc.add_heading('API Request Optimization', level=2)
    doc.add_paragraph('Strategies for optimizing API requests:')
    add_code_snippet(doc, '''
    # Efficient API client management
    @lru_cache()
    def get_openai_client():
        """Cached OpenAI client to avoid repeated initialization"""
        return OpenAI(
            api_key=OPENAI_API_KEY,
            timeout=httpx.Timeout(60.0),
            max_retries=3
        )
    
    # Parallel processing for multiple requests
    async def process_multiple_translations(texts: List[str]):
        tasks = [translate_text(text) for text in texts]
        return await asyncio.gather(*tasks)
    ''')

def add_security_documentation(doc):
    """Add security documentation section"""
    doc.add_heading('6.15 Security Documentation', level=1)
    
    # API Key Management
    doc.add_heading('API Key Management', level=2)
    doc.add_paragraph('Secure API key handling:')
    add_code_snippet(doc, '''
    # Environment-based key management
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
    GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
    
    def validate_api_keys():
        if not OPENAI_API_KEY or not GEMINI_API_KEY:
            raise ValueError("Missing required API keys")
    ''')
    
    # Rate Limiting
    doc.add_heading('Rate Limiting', level=2)
    doc.add_paragraph('Rate limiting implementation:')
    add_code_snippet(doc, '''
    # Rate limiting middleware
    @app.middleware("http")
    async def rate_limit_middleware(request: Request, call_next):
        client_ip = request.client.host
        if await is_rate_limited(client_ip):
            raise HTTPException(status_code=429, detail="Too many requests")
        return await call_next(request)
    ''')
    
    # Input Validation
    doc.add_heading('Input Validation', level=2)
    doc.add_paragraph('Input validation and sanitization:')
    add_code_snippet(doc, '''
    # Input validation
    def validate_input(text: str, max_length: int = 10000) -> bool:
        if not text or len(text) > max_length:
            return False
        return True
    
    # Data sanitization
    def sanitize_input(text: str) -> str:
        return text.strip().replace('<script>', '').replace('</script>', '')
    ''')

def add_deployment_guide(doc):
    """Add deployment guide section"""
    doc.add_heading('6.16 Deployment Guide', level=1)
    
    # Environment Setup
    doc.add_heading('Environment Setup', level=2)
    doc.add_paragraph('Required environment setup:')
    add_code_snippet(doc, '''
    # Environment variables
    OPENAI_API_KEY=your_openai_key
    GEMINI_API_KEY=your_gemini_key
    DATABASE_URL=postgresql://user:password@localhost:5432/dbname
    REDIS_URL=redis://localhost:6379
    ''')
    
    # Deployment Steps
    doc.add_heading('Deployment Steps', level=2)
    doc.add_paragraph('Step-by-step deployment process:')
    add_code_snippet(doc, '''
    # 1. Install dependencies
    pip install -r requirements.txt
    
    # 2. Set up environment
    cp .env.example .env
    # Edit .env with your configuration
    
    # 3. Initialize database
    python manage.py init_db
    
    # 4. Start the application
    uvicorn translation_bot:app --host 0.0.0.0 --port 8000
    ''')
    
    # Monitoring Setup
    doc.add_heading('Monitoring Setup', level=2)
    doc.add_paragraph('Monitoring and logging configuration:')
    add_code_snippet(doc, '''
    # Logging configuration
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('app.log'),
            logging.StreamHandler()
        ]
    )
    
    # Health check endpoint
    @app.get("/health")
    async def health_check():
        return {"status": "healthy"}
    ''')

def add_user_guide(doc):
    """Add user guide section"""
    doc.add_heading('6.17 User Guide', level=1)
    
    # Getting Started
    doc.add_heading('Getting Started', level=2)
    doc.add_paragraph('Quick start guide:')
    add_code_snippet(doc, '''
    # Install dependencies
    pip install -r requirements.txt
    
    # Set up environment variables
    cp .env.example .env
    # Edit .env with your API keys
    
    # Run the application
    uvicorn translation_bot:app --reload
    ''')
    
    # Common Use Cases
    doc.add_heading('Common Use Cases', level=2)
    doc.add_paragraph('Example usage scenarios:')
    add_code_snippet(doc, '''
    # 1. Basic Translation
    curl -X POST "http://localhost:8000/" \\
         -H "Content-Type: application/json" \\
         -d '{"text": "Hello world", "source_language": "en", "target_language": "fa"}'
    
    # 2. Text Editing
    curl -X POST "http://localhost:8000/edit" \\
         -H "Content-Type: application/json" \\
         -d '{"text": "Hello world", "edit_type": "improve"}'
    
    # 3. WebSocket Connection
    websocat ws://localhost:8000/ws
    ''')
    
    # Best Practices
    doc.add_heading('Best Practices', level=2)
    doc.add_paragraph('Recommended usage patterns:')
    add_code_snippet(doc, '''
    # 1. Batch Processing
    async def process_batch(texts: List[str]):
        tasks = [translate_text(text) for text in texts]
        return await asyncio.gather(*tasks)
    
    # 2. Error Handling
    try:
        result = await translate_text(text)
    except TranslationError as e:
        logger.error(f"Translation failed: {e}")
        # Handle error appropriately
    
    # 3. Resource Management
    async with AsyncClient() as client:
        response = await client.post("/", json=data)
    ''')

def add_integration_guide(doc):
    """Add integration guide section"""
    doc.add_heading('6.18 Integration Guide', level=1)
    
    # API Integration
    doc.add_heading('API Integration', level=2)
    doc.add_paragraph('Example API integration:')
    add_code_snippet(doc, '''
    # Python client example
    import requests
    
    def translate_text(text: str, source_lang: str, target_lang: str):
        response = requests.post(
            "http://localhost:8000/",
            json={
                "text": text,
                "source_language": source_lang,
                "target_language": target_lang
            }
        )
        return response.json()
    ''')
    
    # WebSocket Integration
    doc.add_heading('WebSocket Integration', level=2)
    doc.add_paragraph('Real-time updates integration:')
    add_code_snippet(doc, '''
    # WebSocket client example
    import websockets
    
    async def connect_websocket():
        async with websockets.connect('ws://localhost:8000/ws') as websocket:
            while True:
                message = await websocket.recv()
                print(f"Received: {message}")
    ''')
    
    # Third-party Integrations
    doc.add_heading('Third-party Integrations', level=2)
    doc.add_paragraph('Integration with external services:')
    add_code_snippet(doc, '''
    # OpenAI integration
    from openai import AsyncOpenAI
    
    async def openai_translate(text: str):
        client = AsyncOpenAI()
        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a translator."},
                {"role": "user", "content": f"Translate: {text}"}
            ]
        )
        return response.choices[0].message.content
    
    # Gemini integration
    import google.generativeai as genai
    
    def gemini_translate(text: str):
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(f"Translate: {text}")
        return response.text
    ''')

def add_executive_summary(doc):
    """Add executive summary section to the document."""
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'PersianAI is a sophisticated translation system designed to provide accurate '
        'and context-aware translations between Persian and English languages. The system '
        'leverages state-of-the-art AI models from OpenAI and Google, combined with '
        'custom processing logic to handle the unique characteristics of Persian text.'
    )
    doc.add_paragraph(
        'Key features include:\n'
        '• Real-time translation with WebSocket support\n'
        '• Context-aware translation improvements\n'
        '• Detailed change tracking and version control\n'
        '• Comprehensive API for integration\n'
        '• Robust error handling and monitoring'
    )

def add_introduction(doc):
    """Add introduction section to the document."""
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The PersianAI project addresses the growing need for accurate and reliable '
        'Persian-English translation services. By combining advanced AI models with '
        'specialized processing for Persian language characteristics, the system provides '
        'high-quality translations suitable for both casual and professional use.'
    )
    
    doc.add_heading('Project Goals', level=2)
    doc.add_paragraph(
        '• Provide accurate Persian-English translations\n'
        '• Handle complex Persian language structures\n'
        '• Maintain context and meaning in translations\n'
        '• Offer real-time translation capabilities\n'
        '• Support integration with other systems'
    )

def add_system_architecture(doc):
    """Add system architecture section to the document."""
    doc.add_heading('System Architecture', level=1)
    
    doc.add_heading('Overview', level=2)
    doc.add_paragraph(
        'The system follows a microservices architecture with the following components:\n'
        '• FastAPI web server\n'
        '• Translation processing service\n'
        '• Change tracking system\n'
        '• WebSocket communication layer\n'
        '• External API integrations'
    )
    
    doc.add_heading('Component Interaction', level=2)
    doc.add_paragraph(
        'Components interact through well-defined APIs and message queues, ensuring:\n'
        '• Loose coupling between services\n'
        '• Scalability and maintainability\n'
        '• Efficient resource utilization\n'
        '• Reliable error handling'
    )

def add_implementation_details(doc):
    """Add implementation details section to the document."""
    doc.add_heading('Implementation Details', level=1)
    
    doc.add_heading('Core Components', level=2)
    doc.add_paragraph(
        'The implementation consists of several key components:\n'
        '• FastAPI application server\n'
        '• Translation processing engine\n'
        '• Change tracking system\n'
        '• WebSocket server\n'
        '• API integration layer'
    )
    
    doc.add_heading('Technology Stack', level=2)
    doc.add_paragraph(
        'The project utilizes modern technologies:\n'
        '• Python 3.9+\n'
        '• FastAPI framework\n'
        '• OpenAI and Gemini APIs\n'
        '• PostgreSQL database\n'
        '• Redis caching'
    )

def add_testing_and_validation(doc):
    """Add testing and validation section to the document."""
    doc.add_heading('Testing and Validation', level=1)
    
    doc.add_heading('Testing Strategy', level=2)
    doc.add_paragraph(
        'The testing approach includes:\n'
        '• Unit tests for core functionality\n'
        '• Integration tests for API endpoints\n'
        '• Performance testing\n'
        '• Security testing\n'
        '• User acceptance testing'
    )
    
    doc.add_heading('Validation Methods', level=2)
    doc.add_paragraph(
        'Translation quality is validated through:\n'
        '• Automated quality metrics\n'
        '• Manual review by language experts\n'
        '• User feedback analysis\n'
        '• A/B testing of improvements'
    )

def add_api_response_examples(doc):
    """Add API response examples section"""
    doc.add_heading('6.19 API Response Examples', level=1)
    
    # Successful Responses
    doc.add_heading('Successful Responses', level=2)
    doc.add_paragraph('Example successful responses from various endpoints:')
    add_code_snippet(doc, '''
    # Translation endpoint response
    {
        "success": true,
        "translated_text": "Hello world",
        "source_language": "fa",
        "target_language": "en",
        "processing_time": "0.5s",
        "model_used": "gpt-3.5-turbo"
    }
    
    # Edit endpoint response
    {
        "success": true,
        "edited_text": "Improved text version",
        "changes": {
            "additions": 2,
            "deletions": 1,
            "modifications": 3
        },
        "html_diff": "<div class='diff'>...</div>"
    }
    
    # Health check response
    {
        "status": "healthy",
        "api_status": {
            "openai": "connected",
            "gemini": "connected"
        },
        "system_metrics": {
            "memory_usage": "45%",
            "cpu_usage": "30%"
        }
    }
    ''')
    
    # Error Responses
    doc.add_heading('Error Responses', level=2)
    doc.add_paragraph('Example error responses and their meanings:')
    add_code_snippet(doc, '''
    # Authentication error
    {
        "error": "authentication_error",
        "message": "Invalid API key provided",
        "status_code": 401,
        "details": "Please check your API key configuration"
    }
    
    # Rate limit error
    {
        "error": "rate_limit_exceeded",
        "message": "Too many requests",
        "status_code": 429,
        "retry_after": 60
    }
    
    # Validation error
    {
        "error": "validation_error",
        "message": "Invalid input parameters",
        "status_code": 400,
        "details": {
            "text": "Text length exceeds maximum limit",
            "max_length": 10000
        }
    }
    ''')

def add_contribution_guidelines(doc):
    """Add contribution guidelines section"""
    doc.add_heading('6.21 Contribution Guidelines', level=1)
    
    # Code Style
    doc.add_heading('Code Style Guidelines', level=2)
    doc.add_paragraph('Follow these style guidelines when contributing:')
    add_code_snippet(doc, '''
    # Python Style Guide
    1. Follow PEP 8 conventions
    2. Use meaningful variable names
    3. Add docstrings to all functions
    4. Keep functions focused and small
    5. Write comprehensive tests
    
    # Documentation Style
    1. Keep documentation up to date
    2. Use clear and concise language
    3. Include code examples
    4. Document all parameters
    5. Explain return values
    ''')
    
    # Pull Request Process
    doc.add_heading('Pull Request Process', level=2)
    doc.add_paragraph('Steps for submitting changes:')
    add_code_snippet(doc, '''
    1. Fork the repository
    2. Create a feature branch
    3. Make your changes
    4. Write/update tests
    5. Update documentation
    6. Submit pull request
    
    # PR Template
    ## Description
    Brief description of changes
    
    ## Changes Made
    - Detailed list of changes
    - Impact on existing features
    - New features added
    
    ## Testing
    - Tests added/updated
    - Test coverage report
    ''')
    
    # Issue Reporting
    doc.add_heading('Issue Reporting', level=2)
    doc.add_paragraph('Guidelines for reporting issues:')
    add_code_snippet(doc, '''
    # Issue Template
    ## Problem Description
    Detailed description of the issue
    
    ## Steps to Reproduce
    1. Step-by-step reproduction
    2. Expected behavior
    3. Actual behavior
    
    ## Environment
    - OS version
    - Python version
    - Dependencies
    ''')

def add_changelog(doc):
    """Add changelog section"""
    doc.add_heading('6.22 Changelog', level=1)
    
    # Version History
    doc.add_heading('Version History', level=2)
    doc.add_paragraph('Recent changes and updates:')
    add_code_snippet(doc, '''
    # Version 1.2.0 (2024-03-15)
    ## Added
    - Support for Gemini Pro model
    - Real-time translation updates
    - Enhanced error handling
    
    ## Changed
    - Improved translation accuracy
    - Updated API response format
    - Optimized caching system
    
    ## Fixed
    - Memory leak in WebSocket connections
    - Rate limiting issues
    - Database connection handling
    
    # Version 1.1.0 (2024-02-01)
    ## Added
    - WebSocket support
    - Translation history
    - User preferences
    
    ## Changed
    - Updated OpenAI integration
    - Enhanced logging system
    - Improved error messages
    ''')
    
    # Breaking Changes
    doc.add_heading('Breaking Changes', level=2)
    doc.add_paragraph('Notable breaking changes and migration guides:')
    add_code_snippet(doc, '''
    # Version 1.2.0
    ## Breaking Changes
    1. API response format updated
    2. New authentication requirements
    3. Changed database schema
    
    ## Migration Guide
    1. Update API clients to handle new response format
    2. Generate new API keys
    3. Run database migration script
    
    # Version 1.1.0
    ## Breaking Changes
    1. Deprecated old REST endpoints
    2. Changed configuration format
    
    ## Migration Guide
    1. Update to new WebSocket endpoints
    2. Update configuration files
    ''')

def add_system_requirements(doc):
    """Add system requirements section"""
    doc.add_heading('6.23 System Requirements', level=1)
    
    # Hardware Requirements
    doc.add_heading('Hardware Requirements', level=2)
    doc.add_paragraph('Minimum hardware specifications:')
    add_code_snippet(doc, '''
    # Production Environment
    - CPU: 4+ cores
    - RAM: 8GB minimum, 16GB recommended
    - Storage: 20GB minimum
    - Network: 100Mbps minimum
    
    # Development Environment
    - CPU: 2+ cores
    - RAM: 4GB minimum
    - Storage: 10GB minimum
    - Network: 50Mbps minimum
    ''')
    
    # Software Dependencies
    doc.add_heading('Software Dependencies', level=2)
    doc.add_paragraph('Required software and versions:')
    add_code_snippet(doc, '''
    # Core Dependencies
    - Python 3.9+
    - PostgreSQL 13+
    - Redis 6+
    - Node.js 16+ (for frontend)
    
    # Python Packages
    - fastapi==0.68.1
    - uvicorn==0.15.0
    - python-dotenv==0.19.0
    - openai==1.0.0
    - google-generativeai==0.3.0
    
    # System Software
    - Docker 20+
    - nginx 1.18+
    - Supervisor 4+
    ''')
    
    # Operating System Compatibility
    doc.add_heading('Operating System Compatibility', level=2)
    doc.add_paragraph('Supported operating systems and configurations:')
    add_code_snippet(doc, '''
    # Linux
    - Ubuntu 20.04 LTS or newer
    - CentOS 8 or newer
    - Debian 11 or newer
    
    # macOS
    - macOS 11 (Big Sur) or newer
    - Apple Silicon supported
    
    # Windows
    - Windows 10 or newer
    - Windows Server 2019 or newer
    ''')
    
    # Network Requirements
    doc.add_heading('Network Requirements', level=2)
    doc.add_paragraph('Network configuration requirements:')
    add_code_snippet(doc, '''
    # Firewall Rules
    - Port 80/443 for HTTP/HTTPS
    - Port 5432 for PostgreSQL
    - Port 6379 for Redis
    - Port 8000 for API server
    
    # External Services
    - OpenAI API access
    - Google Cloud API access
    - CDN endpoints
    - Monitoring services
    ''')

def create_system_architecture_diagram(doc):
    """Create and add system architecture diagram to the document."""
    doc.add_heading("System Architecture Diagram", level=3)
    
    # Create a table to represent the architecture diagram
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(2)
    
    # Frontend Layer
    table.cell(0, 1).text = "Frontend Web Interface"
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Backend Layer
    table.cell(1, 0).text = "FastAPI Backend Server"
    table.cell(1, 1).text = "Translation Service"
    table.cell(1, 2).text = "WebSocket Server"
    
    # External Services Layer
    table.cell(2, 0).text = "OpenAI API"
    table.cell(2, 1).text = "Gemini API"
    table.cell(2, 2).text = "Database/Cache"
    
    # Add arrows and connections
    doc.add_paragraph("↑ API Requests")
    doc.add_paragraph("↑ Real-time Updates")
    doc.add_paragraph("↑ Data Storage")
    
    # Add description
    doc.add_paragraph(
        "The system architecture consists of three main layers:\n"
        "1. Frontend Layer: Web interface for user interaction\n"
        "2. Backend Layer: Core translation and processing services\n"
        "3. External Services Layer: AI models and data storage"
    )

def create_data_flow_diagram(doc):
    """Create and add data flow diagram to the document."""
    doc.add_heading("Data Flow Diagram", level=3)
    
    # Create a table to represent the data flow
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    
    # Set row heights
    for row in table.rows:
        for cell in row.cells:
            cell.height = Inches(0.5)
    
    # Add flow steps
    table.cell(0, 0).text = "1. User Input → Frontend Validation"
    table.cell(1, 0).text = "2. Frontend → Backend API Request"
    table.cell(2, 0).text = "3. Backend → AI Model Processing"
    table.cell(3, 0).text = "4. AI Model → Response Formatting"
    table.cell(4, 0).text = "5. Response → User Interface"
    
    # Add arrows between steps
    for i in range(4):
        doc.add_paragraph("↓")
        doc.add_paragraph("Data Flow")
    
    # Add description
    doc.add_paragraph(
        "The data flow diagram shows the sequence of data processing:\n"
        "• Input validation ensures data quality\n"
        "• API requests are authenticated and rate-limited\n"
        "• AI models process text in chunks\n"
        "• Responses are formatted and validated\n"
        "• Results are delivered to the user interface"
    )

def create_error_handling_flowchart(doc):
    """Create and add error handling flowchart to the document."""
    doc.add_heading("Error Handling Flowchart", level=3)
    
    # Create a table to represent the flowchart
    table = doc.add_table(rows=6, cols=1)
    table.style = 'Table Grid'
    
    # Set row heights
    for row in table.rows:
        for cell in row.cells:
            cell.height = Inches(0.5)
    
    # Add error handling steps
    table.cell(0, 0).text = "1. Error Detection"
    table.cell(1, 0).text = "2. Error Classification"
    table.cell(2, 0).text = "3. Error Handling Strategy"
    table.cell(3, 0).text = "4. Recovery Action"
    table.cell(4, 0).text = "5. User Notification"
    table.cell(5, 0).text = "6. Logging and Monitoring"
    
    # Add arrows between steps
    for i in range(5):
        doc.add_paragraph("↓")
        doc.add_paragraph("Error Handling Flow")
    
    # Add description
    doc.add_paragraph(
        "The error handling flowchart illustrates the system's error management:\n"
        "• Errors are detected through validation and monitoring\n"
        "• Errors are classified by type and severity\n"
        "• Appropriate handling strategies are selected\n"
        "• Recovery actions are executed when possible\n"
        "• Users are notified of errors and their status\n"
        "• All errors are logged for analysis and improvement"
    )

def add_visual_aids(doc):
    """Add visual aids section to the document."""
    doc.add_heading("6.13 Visual Aids and Diagrams", level=2)
    
    # Add system architecture diagram
    create_system_architecture_diagram(doc)
    
    # Add data flow diagram
    create_data_flow_diagram(doc)
    
    # Add error handling flowchart
    create_error_handling_flowchart(doc)

def add_real_world_examples(doc):
    """Add real-world examples section to the document."""
    doc.add_heading("6.14 Real-World Examples and Use Cases", level=2)
    
    # Business Document Translation
    doc.add_heading("Business Document Translation", level=3)
    doc.add_paragraph(
        "Example of translating a business contract from Persian to English:"
    )
    add_code_snippet(
        doc,
        """# Example API request
POST /translate
{
    "text": "قرارداد همکاری تجاری بین شرکت الف و شرکت ب...",
    "source_lang": "fa",
    "target_lang": "en",
    "format": "business"
}

# Example response
{
    "translated_text": "Commercial Cooperation Agreement between Company A and Company B...",
    "confidence_score": 0.95,
    "processing_time": 1.2
}"""
    )

    # Academic Paper Translation
    doc.add_heading("Academic Paper Translation", level=3)
    doc.add_paragraph(
        "Example of translating an academic paper with technical terminology:"
    )
    add_code_snippet(
        doc,
        """# Example API request
POST /translate
{
    "text": "مقاله علمی در زمینه هوش مصنوعی و یادگیری ماشین...",
    "source_lang": "fa",
    "target_lang": "en",
    "format": "academic",
    "domain": "computer_science"
}"""
    )

    # Website Content Translation
    doc.add_heading("Website Content Translation", level=3)
    doc.add_paragraph(
        "Example of translating website content with HTML preservation:"
    )
    add_code_snippet(
        doc,
        """# Example API request
POST /translate
{
    "text": "<h1>خوش آمدید</h1><p>به وبسایت ما خوش آمدید...</p>",
    "source_lang": "fa",
    "target_lang": "en",
    "format": "html",
    "preserve_tags": true
}"""
    )

def add_development_workflow(doc):
    """Add development workflow section to the document."""
    doc.add_heading("6.15 Development Workflow and CI/CD", level=2)
    
    # Development Environment Setup
    doc.add_heading("Development Environment Setup", level=3)
    doc.add_paragraph(
        "Steps to set up the development environment:"
    )
    add_code_snippet(
        doc,
        """# Clone the repository
git clone https://github.com/your-org/persian-ai.git

# Create and activate virtual environment
python -m venv .venv
source .venv/bin/activate  # On Windows: .venv\\Scripts\\activate

# Install dependencies
pip install -r requirements.txt

# Set up environment variables
cp .env.example .env
# Edit .env with your API keys and configuration

# Run tests
pytest tests/

# Start development server
uvicorn translation_bot:app --reload"""
    )

    # CI/CD Pipeline
    doc.add_heading("CI/CD Pipeline", level=3)
    doc.add_paragraph(
        "The project uses GitHub Actions for continuous integration and deployment:"
    )
    doc.add_paragraph(
        "1. Code Quality Checks:\n"
        "- Linting with flake8\n"
        "- Type checking with mypy\n"
        "- Code formatting with black\n"
        "\n"
        "2. Testing:\n"
        "- Unit tests with pytest\n"
        "- Integration tests\n"
        "- Performance tests\n"
        "\n"
        "3. Deployment:\n"
        "- Automated deployment to staging\n"
        "- Production deployment with approval\n"
        "- Version tagging and release notes"
    )

    # Git Workflow
    doc.add_heading("Git Workflow", level=3)
    doc.add_paragraph(
        "The project follows a feature branch workflow:"
    )
    doc.add_paragraph(
        "1. Feature Development:\n"
        "- Create feature branch from main\n"
        "- Develop and test locally\n"
        "- Push changes and create PR\n"
        "\n"
        "2. Code Review:\n"
        "- Peer review required\n"
        "- CI checks must pass\n"
        "- Address review comments\n"
        "\n"
        "3. Merging:\n"
        "- Squash and merge to main\n"
        "- Delete feature branch\n"
        "- Update version if needed"
    )

def add_model_limitations(doc):
    """Add model limitations section to the document."""
    doc.add_heading("6.19 Model Limitations", level=1)
    
    # OpenAI Models
    doc.add_heading("OpenAI Models", level=2)
    
    # GPT-3.5-turbo
    doc.add_heading("GPT-3.5-turbo", level=3)
    doc.add_paragraph(
        "Key limitations and specifications:\n"
        "• Maximum tokens: 30,000\n"
        "• Maximum chunk size: 1,500 characters\n"
        "• Maximum text length: 15,000 characters (approximately 3000 words)\n"
        "• Maximum output tokens: 4,000\n"
        "• Best for: Fast and reliable processing with good accuracy\n"
        "• Estimated processing time for 3000 words: 2-3 minutes"
    )
    
    # GPT-4
    doc.add_heading("GPT-4", level=3)
    doc.add_paragraph(
        "Key limitations and specifications:\n"
        "• Maximum tokens: 50,000\n"
        "• Maximum chunk size: 2,500 characters\n"
        "• Maximum text length: 21,000 characters (approximately 3000 words)\n"
        "• Maximum output tokens: 4,000\n"
        "• Best for: Most accurate processing, better understanding of context and nuances\n"
        "• Estimated processing time for 3000 words: 1.5-2 minutes"
    )
    
    # Gemini Models
    doc.add_heading("Gemini Models", level=2)
    
    # Gemini 1.5 Flash
    doc.add_heading("Gemini 1.5 Flash (gemini-1.5-flash-8b)", level=3)
    doc.add_paragraph(
        "Key limitations and specifications:\n"
        "• Maximum tokens: 1,000,000\n"
        "• Maximum output tokens: 1,024\n"
        "• Maximum text length: 15,000 characters (approximately 3000 words)\n"
        "• Best for: Fast and efficient processing with good accuracy\n"
        "• Estimated processing time for 3000 words: 1-1.5 minutes"
    )
    
    # Gemini 1.5 Pro
    doc.add_heading("Gemini 1.5 Pro (gemini-1.5-pro-latest)", level=3)
    doc.add_paragraph(
        "Key limitations and specifications:\n"
        "• Maximum tokens: 2,000,000\n"
        "• Maximum output tokens: 1,024\n"
        "• Maximum text length: 21,000 characters (approximately 3000 words)\n"
        "• Best for: Advanced AI model with strong understanding of Persian language\n"
        "• Estimated processing time for 3000 words: 1-1.5 minutes"
    )
    
    # General Notes
    doc.add_heading("General Notes", level=2)
    doc.add_paragraph(
        "Important considerations for all models:\n"
        "• All models have built-in chunking mechanisms to handle long texts\n"
        "• Word count validation ensures edited text stays within 20 words of the original\n"
        "• Timeout limits are set to 60 seconds for translation requests\n"
        "• Each model has retry mechanisms for failed requests\n"
        "• For texts longer than 1000 words, Gemini models are recommended for better performance\n"
        "• Processing times are estimates and may vary based on text complexity and API response times"
    )

def add_environment_setup_documentation(doc):
    """Add environment setup and API key configuration documentation."""
    doc.add_heading('6.15 Environment Setup and API Keys', level=1)
    
    # Environment Variables
    doc.add_heading('Environment Variables', level=2)
    p = doc.add_paragraph('The system uses environment variables for configuration and API keys. Create a `.env` file in the root directory with the following variables:')
    
    # Create a table for environment variables
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Variable Name'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Example Value'
    
    # Add environment variables
    env_vars = [
        ('OPENAI_API_KEY', 'Your OpenAI API key for GPT-4 access', 'sk-...'),
        ('GEMINI_API_KEY', 'Your Google Gemini API key', 'AIza...'),
        ('DATABASE_URL', 'PostgreSQL database connection URL', 'postgresql://user:pass@localhost:5432/dbname'),
        ('REDIS_URL', 'Redis connection URL for caching', 'redis://localhost:6379/0'),
        ('LOG_LEVEL', 'Logging level (DEBUG, INFO, WARNING, ERROR)', 'INFO'),
        ('MAX_TOKENS', 'Maximum tokens for API responses', '2000'),
        ('CACHE_TTL', 'Cache time-to-live in seconds', '3600'),
        ('RATE_LIMIT', 'API rate limit per minute', '60'),
        ('ENVIRONMENT', 'Deployment environment (development/production)', 'development')
    ]
    
    for var_name, description, example in env_vars:
        row_cells = table.add_row().cells
        row_cells[0].text = var_name
        row_cells[1].text = description
        row_cells[2].text = example
    
    # API Key Setup Instructions
    doc.add_heading('API Key Setup Instructions', level=2)
    
    # OpenAI API Key
    doc.add_heading('OpenAI API Key', level=3)
    p = doc.add_paragraph('To obtain an OpenAI API key:')
    steps = [
        '1. Visit https://platform.openai.com/',
        '2. Create an account or sign in',
        '3. Navigate to API Keys section',
        '4. Create a new API key',
        '5. Copy the key and add it to your .env file',
        '6. Never share or commit your API key'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Gemini API Key
    doc.add_heading('Gemini API Key', level=3)
    p = doc.add_paragraph('To obtain a Gemini API key:')
    steps = [
        '1. Visit https://makersuite.google.com/app/apikey',
        '2. Sign in with your Google account',
        '3. Create a new API key',
        '4. Copy the key and add it to your .env file',
        '5. Keep your API key secure'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Security Best Practices
    doc.add_heading('Security Best Practices', level=2)
    practices = [
        'Never commit the .env file to version control',
        'Use different API keys for development and production',
        'Regularly rotate API keys',
        'Set up API key usage monitoring',
        'Implement rate limiting',
        'Use environment-specific .env files (.env.development, .env.production)'
    ]
    for practice in practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    # Example .env file
    doc.add_heading('Example .env File', level=2)
    p = doc.add_paragraph('Example .env file structure:')
    p = doc.add_paragraph('''# API Keys
OPENAI_API_KEY=sk-your-openai-key-here
GEMINI_API_KEY=AIza-your-gemini-key-here

# Database Configuration
DATABASE_URL=postgresql://user:password@localhost:5432/dbname
REDIS_URL=redis://localhost:6379/0

# Application Settings
LOG_LEVEL=INFO
MAX_TOKENS=2000
CACHE_TTL=3600
RATE_LIMIT=60
ENVIRONMENT=development''')
    p.style = 'Code'

def add_detailed_module_documentation(doc):
    """Add detailed documentation for changes.py and improvements.py modules."""
    doc.add_heading('6.24 Module Documentation', level=1)
    
    # changes.py documentation
    doc.add_heading('changes.py Module', level=2)
    doc.add_paragraph('The changes.py module handles version control, change tracking, and text comparison functionality.')
    
    # Core Functions in changes.py
    doc.add_heading('Core Functions', level=3)
    
    # Track Changes Function
    doc.add_paragraph('track_changes() Function:')
    doc.add_paragraph('Purpose: Tracks and records changes between original and modified text.')
    add_code_snippet(doc, '''
    def track_changes(original: str, modified: str) -> Dict[str, Any]:
        """Track changes between original and modified text.
        
        Args:
            original (str): Original text
            modified (str): Modified text
            
        Returns:
            Dict containing:
            - replacements: Number of replaced segments
            - deletions: Number of deleted segments
            - insertions: Number of inserted segments
            - details: List of specific changes
        """
        changes = {
            'replacements': 0,
            'deletions': 0,
            'insertions': 0,
            'details': []
        }
        # Implementation details...
    ''')
    
    # Generate HTML Diff
    doc.add_paragraph('generate_html_diff() Function:')
    doc.add_paragraph('Purpose: Generates HTML visualization of text differences.')
    add_code_snippet(doc, '''
    def generate_html_diff(original: str, modified: str) -> str:
        """Generate HTML visualization of differences.
        
        Args:
            original (str): Original text
            modified (str): Modified text
            
        Returns:
            str: HTML string showing differences with color coding
        """
        # Implementation details...
    ''')
    
    # Version Control
    doc.add_paragraph('Version Control Functions:')
    doc.add_paragraph('The module includes comprehensive version control:')
    add_code_snippet(doc, '''
    class VersionControl:
        def save_version(self, text: str, version: str) -> None:
            """Save a new version of the text."""
            
        def get_version(self, version: str) -> str:
            """Retrieve a specific version."""
            
        def list_versions(self) -> List[str]:
            """List all available versions."""
            
        def compare_versions(self, v1: str, v2: str) -> Dict[str, Any]:
            """Compare two versions and return differences."""
    ''')
    
    # improvements.py documentation
    doc.add_heading('improvements.py Module', level=2)
    doc.add_paragraph('The improvements.py module contains logic for enhancing translation quality.')
    
    # Translation Prompts
    doc.add_heading('Translation Prompts', level=3)
    doc.add_paragraph('Specialized prompts for different translation scenarios:')
    add_code_snippet(doc, '''
    # Base translation prompt
    TRANSLATION_PROMPT = """
    Translate the following Persian text to English:
    - Maintain formal/informal tone
    - Preserve cultural context
    - Keep technical terms accurate
    - Maintain formatting and structure
    {text}
    """
    
    # Technical translation prompt
    TECHNICAL_TRANSLATION_PROMPT = """
    Translate the following technical Persian text:
    - Preserve technical terminology
    - Maintain academic/technical tone
    - Keep formatting and citations
    - Include glossary for key terms
    {text}
    """
    
    # Literary translation prompt
    LITERARY_TRANSLATION_PROMPT = """
    Translate the following literary Persian text:
    - Preserve literary style and tone
    - Maintain metaphors and cultural references
    - Keep poetic elements where applicable
    - Preserve author's voice
    {text}
    """
    ''')
    
    # Enhancement Functions
    doc.add_heading('Enhancement Functions', level=3)
    doc.add_paragraph('Functions for improving translation quality:')
    
    # Context Analysis
    doc.add_paragraph('analyze_context() Function:')
    doc.add_paragraph('Purpose: Analyzes text context for better translation.')
    add_code_snippet(doc, '''
    def analyze_context(text: str) -> Dict[str, Any]:
        """Analyze text context for better translation.
        
        Args:
            text (str): Input text
            
        Returns:
            Dict containing:
            - domain: Technical domain (e.g., medical, legal)
            - formality: Text formality level
            - style: Writing style
            - key_terms: Important terminology
        """
        # Implementation details...
    ''')
    
    # Quality Enhancement
    doc.add_paragraph('enhance_translation() Function:')
    doc.add_paragraph('Purpose: Applies quality improvements to translations.')
    add_code_snippet(doc, '''
    async def enhance_translation(
        text: str,
        context: Dict[str, Any],
        model_type: str = "gpt-4"
    ) -> str:
        """Enhance translation quality.
        
        Args:
            text (str): Translated text
            context (dict): Context information
            model_type (str): AI model to use
            
        Returns:
            str: Enhanced translation
        """
        # Apply Persian-specific rules
        text = apply_persian_rules(text)
        
        # Generate context-aware prompt
        prompt = generate_context_prompt(text, context)
        
        # Get model suggestions
        suggestions = await get_model_suggestions(prompt)
        
        # Apply improvements
        improved_text = apply_improvements(text, suggestions)
        
        # Validate quality
        quality_score = calculate_quality_score(improved_text)
        
        if quality_score < QUALITY_THRESHOLD:
            improved_text = await apply_additional_improvements(improved_text)
        
        return improved_text
    ''')
    
    # Persian Language Rules
    doc.add_heading('Persian Language Rules', level=3)
    doc.add_paragraph('Specialized rules for Persian language processing:')
    add_code_snippet(doc, '''
    class PersianRules:
        @staticmethod
        def fix_spacing(text: str) -> str:
            """Fix Persian text spacing issues."""
            
        @staticmethod
        def normalize_characters(text: str) -> str:
            """Normalize Persian characters."""
            
        @staticmethod
        def fix_punctuation(text: str) -> str:
            """Fix Persian punctuation."""
            
        @staticmethod
        def handle_numbers(text: str) -> str:
            """Handle Persian numbers and dates."""
    ''')
    
    # Quality Metrics
    doc.add_heading('Quality Metrics', level=3)
    doc.add_paragraph('Functions for measuring translation quality:')
    add_code_snippet(doc, '''
    def calculate_quality_score(text: str) -> float:
        """Calculate translation quality score.
        
        Metrics include:
        - Grammar correctness
        - Terminology accuracy
        - Style consistency
        - Cultural appropriateness
        - Technical accuracy
        """
        
    def validate_technical_terms(text: str, domain: str) -> bool:
        """Validate technical terminology."""
        
    def check_style_consistency(text: str, style: str) -> bool:
        """Check for consistent writing style."""
    ''')

def add_project_files_documentation(doc):
    """Add documentation for all project files and their purposes."""
    doc.add_heading('6.25 Project Files Documentation', level=1)
    
    # Core Application Files
    doc.add_heading('Core Application Files', level=2)
    
    # translation_bot.py
    doc.add_heading('translation_bot.py', level=3)
    doc.add_paragraph('Main application file that handles the core translation functionality.')
    add_code_snippet(doc, '''
    # Key components:
    - FastAPI application setup
    - WebSocket connections
    - Translation endpoints
    - AI model integration
    - Error handling
    - Rate limiting
    - Caching system
    ''')
    
    # changes.py
    doc.add_heading('changes.py', level=3)
    doc.add_paragraph('Handles version control and change tracking for translations.')
    add_code_snippet(doc, '''
    # Key features:
    - Version history management
    - Change tracking
    - HTML diff generation
    - Version comparison
    - Change statistics
    ''')
    
    # improvements.py
    doc.add_heading('improvements.py', level=3)
    doc.add_paragraph('Enhances translation quality and handles Persian-specific rules.')
    add_code_snippet(doc, '''
    # Key features:
    - Context analysis
    - Quality enhancement
    - Persian language rules
    - Technical term handling
    - Style consistency
    ''')
    
    # Test Files
    doc.add_heading('Test Files', level=2)
    
    # test_changes.py
    doc.add_heading('test_changes.py', level=3)
    doc.add_paragraph('Tests for the changes module functionality.')
    add_code_snippet(doc, '''
    # Test cases:
    - Version tracking
    - Change detection
    - HTML diff generation
    - Version comparison
    - Change statistics
    ''')
    
    # test_connections.py
    doc.add_heading('test_connections.py', level=3)
    doc.add_paragraph('Tests API connections and external service integration.')
    add_code_snippet(doc, '''
    # Test cases:
    - OpenAI API connection
    - Gemini API connection
    - Database connection
    - Redis connection
    - WebSocket connection
    ''')
    
    # test_gemini.py
    doc.add_heading('test_gemini.py', level=3)
    doc.add_paragraph('Specific tests for Gemini model integration.')
    add_code_snippet(doc, '''
    # Test cases:
    - API key validation
    - Model response
    - Error handling
    - Rate limiting
    - Response formatting
    ''')
    
    # Configuration Files
    doc.add_heading('Configuration Files', level=2)
    
    # .env
    doc.add_heading('.env', level=3)
    doc.add_paragraph('Environment configuration file for API keys and settings.')
    add_code_snippet(doc, '''
    # Configuration items:
    - API keys
    - Database URLs
    - Cache settings
    - Log levels
    - Rate limits
    - Environment type
    ''')
    
    # requirements.txt
    doc.add_heading('requirements.txt', level=3)
    doc.add_paragraph('Python package dependencies for the project.')
    add_code_snippet(doc, '''
    # Key dependencies:
    - fastapi
    - openai
    - google-generativeai
    - python-dotenv
    - asyncpg
    - redis
    - python-docx
    ''')
    
    # HTML Templates
    doc.add_heading('HTML Templates', level=2)
    
    # index.html
    doc.add_heading('index.html', level=3)
    doc.add_paragraph('Main web interface for the translation system.')
    add_code_snippet(doc, '''
    # Features:
    - Translation interface
    - Real-time updates
    - Error display
    - Version history
    - Change visualization
    ''')
    
    # Static Files
    doc.add_heading('Static Files', level=2)
    
    # CSS
    doc.add_heading('CSS Files', level=3)
    doc.add_paragraph('Styling for the web interface.')
    add_code_snippet(doc, '''
    # Styles:
    - Main layout
    - Translation interface
    - Error messages
    - Version history
    - Change highlights
    ''')
    
    # JavaScript
    doc.add_heading('JavaScript Files', level=3)
    doc.add_paragraph('Client-side functionality for the web interface.')
    add_code_snippet(doc, '''
    # Features:
    - WebSocket connection
    - Real-time updates
    - Form handling
    - Error handling
    - Version management
    ''')
    
    # File Relationships
    doc.add_heading('File Relationships', level=2)
    doc.add_paragraph('How the files work together:')
    
    relationships = [
        ('translation_bot.py', 'Main application that coordinates all components'),
        ('changes.py', 'Tracks and manages translation versions'),
        ('improvements.py', 'Enhances translation quality'),
        ('test_*.py', 'Ensures functionality of each component'),
        ('.env', 'Provides configuration for all components'),
        ('requirements.txt', 'Lists dependencies for all components'),
        ('HTML/CSS/JS', 'Provides user interface for the system')
    ]
    
    for file, purpose in relationships:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{file}: ').bold = True
        p.add_run(purpose)

def generate_report():
    """Generate the technical report."""
    doc = Document()
    
    # Create custom styles
    create_title_style(doc)
    create_heading_style(doc, 1, 14)
    create_heading_style(doc, 2, 12)
    create_heading_style(doc, 3, 11)
    create_normal_style(doc)
    create_code_style(doc)
    
    # Add title page
    add_title_page(doc, 
                  title="PersianAI: Persian-English Translation System",
                  author="PersianAI Team",
                  date=datetime.now().strftime("%B %d, %Y"))
    
    # Add existing sections
    add_executive_summary(doc)
    add_introduction(doc)
    add_system_architecture(doc)
    add_implementation_details(doc)
    add_testing_and_validation(doc)
    add_code_documentation(doc)
    add_function_documentation(doc)
    add_api_endpoints_documentation(doc)
    add_error_handling_documentation(doc)
    add_performance_optimization_documentation(doc)
    add_security_documentation(doc)
    add_deployment_guide(doc)
    add_user_guide(doc)
    add_integration_guide(doc)
    add_visual_aids(doc)
    add_environment_setup_documentation(doc)
    add_detailed_module_documentation(doc)
    add_project_files_documentation(doc)  # Add the new project files documentation
    
    # Save the document
    doc.save('PersianAI_Technical_Report.docx')
    print("Technical report generated successfully as 'PersianAI_Technical_Report.docx'")

if __name__ == "__main__":
    generate_report() 