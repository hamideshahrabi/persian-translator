from fastapi import FastAPI, HTTPException, Request
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Union, Tuple, Any
import os
from dotenv import load_dotenv
import logging
import json
from enum import Enum
import google.generativeai as genai
import difflib
import re
import openai
import asyncio
from functools import lru_cache
from difflib import SequenceMatcher
from html import escape
from contextlib import asynccontextmanager
import socket
import httpx
import aiohttp
import backoff  # For exponential backoff
from typing import Optional, Dict, Any
import threading
import time
from dataclasses import dataclass
from datetime import datetime, timedelta
from improvements import TRANSLATION_PROMPT, EDIT_PROMPT, EDIT_PROMPT_DETAILED
import uvicorn
from fastapi.middleware.cors import CORSMiddleware
import psutil
from openai import OpenAI, AsyncOpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_COLOR
from io import BytesIO
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import tempfile

# Configure logging and load API keys
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables once at startup
load_dotenv(override=True)  # Force reload of environment variables

# Initialize API keys
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable is not set")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable is not set")

# Initialize Gemini
genai.configure(api_key=GEMINI_API_KEY)

# Initialize OpenAI with better configuration
@lru_cache()
def get_sync_openai_client():
    """Get or create synchronous OpenAI client instance."""
    try:
        return OpenAI(
            api_key=OPENAI_API_KEY,
            timeout=httpx.Timeout(60.0, connect=10.0, read=30.0, write=30.0),
            max_retries=3
        )
    except Exception as e:
        logger.error(f"Error initializing sync OpenAI client: {str(e)}")
        raise

@lru_cache()
def get_async_openai_client():
    """Get an async OpenAI client with proper configuration."""
    client = AsyncOpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        timeout=60.0,  # Increase timeout to 60 seconds
        max_retries=3  # Add retries for better reliability
    )
    return client

# Initialize models
@lru_cache()
def get_gemini_model(model_name: str = "models/gemini-1.5-pro-latest") -> genai.GenerativeModel:
    """Get or create Gemini model instance."""
    try:
        generation_config = {
            "temperature": 0.1,  # Lower temperature for faster, more deterministic responses
            "top_p": 0.95,      # Higher top_p for faster sampling
            "top_k": 20,        # Lower top_k for faster sampling
            "max_output_tokens": 2048,  # Increased for longer texts
            "candidate_count": 1,  # Only generate one candidate
        }
        
        logger.info(f"Initializing Gemini model: {model_name}")
        model = genai.GenerativeModel(
            model_name=model_name,
            generation_config=generation_config
        )
        
        # Test the model with a simple prompt
        logger.info("Testing Gemini model with simple prompt...")
        response = model.generate_content("Test.")
        if not response or not response.text:
            logger.error("Failed to generate content with Gemini model")
            raise Exception("Failed to generate content with Gemini model")
            
        logger.info("Gemini model initialized successfully")
        return model
    except Exception as e:
        logger.error(f"Error initializing Gemini model: {str(e)}")
        raise

# Initialize FastAPI app
app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:9090", "http://localhost:9091", "http://localhost:9092"],  # Add your production domains here
    allow_credentials=True,
    allow_methods=["GET", "POST"],  # Specify only the methods you need
    allow_headers=["Content-Type", "Authorization"],  # Specify only the headers you need
)

# Setup templates
templates = Jinja2Templates(directory="templates")

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Define editing stages and context
EDIT_STAGES = {
    'grammar': {
        'name': 'اصلاح دستور زبان و نگارش',
        'focus': [
            'اصلاح خطاهای دستوری',
            'بهبود نقطه‌گذاری',
            'اصلاح فاصله‌گذاری'
        ],
        'examples': 'مثال: تبدیل "من رفتم خانه." به "من به خانه رفتم."'
    },
    'clarity': {
        'name': 'بهبود وضوح و روانی',
        'focus': [
            'بهبود ساختار جملات',
            'حذف ابهام',
            'افزایش روانی متن'
        ],
        'examples': 'مثال: ساده‌سازی جملات پیچیده با حفظ معنی'
    },
    'professional': {
        'name': 'استانداردسازی حرفه‌ای',
        'focus': [
            'یکدست‌سازی اصطلاحات',
            'حفظ لحن حرفه‌ای',
            'بهبود انسجام متن'
        ],
        'examples': 'مثال: استفاده از اصطلاحات استاندارد حرفه‌ای'
    }
}

# Global variables for API clients and keys
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

# Connection settings
MAX_RETRIES = 3
TIMEOUT = 30.0
BACKOFF_FACTOR = 2
INITIAL_BACKOFF = 1

# Validate API keys with proper error handling
def validate_api_keys():
    """Validate API keys and raise descriptive errors."""
    missing_keys = []
    if not OPENAI_API_KEY:
        missing_keys.append("OpenAI API key")
    if not GEMINI_API_KEY:
        missing_keys.append("Gemini API key")
    
    if missing_keys:
        raise ValueError(f"Missing required API keys: {', '.join(missing_keys)}")

# Initialize API clients with proper error handling
async def initialize_api_clients():
    """Initialize API clients and verify they work."""
    try:
        # Test OpenAI client
        client = get_sync_openai_client()
        models = client.models.list()  # Sync call for initialization
        logger.info("OpenAI client initialized successfully")
        logger.info(f"Available OpenAI models: {[model.id for model in models.data]}")

        # Test Gemini model
        model = get_gemini_model()
        logger.info("Gemini model initialized successfully")
        
        return True
    except Exception as e:
        logger.error(f"Failed to initialize API clients: {str(e)}")
        return False

class ModelType(str, Enum):
    GPT35 = "gpt-3.5-turbo"
    GPT4 = "gpt-4"
    GEMINI_FLASH = "models/gemini-1.5-flash-8b"
    GEMINI_PRO = "models/gemini-1.5-pro-latest"
    
    @property
    def description(self) -> str:
        descriptions = {
            self.GPT35: "Fast and reliable processing with good accuracy",
            self.GPT4: "Most accurate processing, better understanding of context and nuances",
            self.GEMINI_FLASH: "Gemini 1.5 Flash - Fast and efficient model for text processing",
            self.GEMINI_PRO: "Gemini 1.5 Pro - Advanced model with better understanding"
        }
        return descriptions.get(self, "Unknown model")
        
    @property
    def max_tokens(self) -> int:
        limits = {
            self.GPT35: 30000,
            self.GPT4: 50000,
            self.GEMINI_FLASH: 1_000_000,
            self.GEMINI_PRO: 2000000
        }
        return limits.get(self, 30000)  # Default to 30000 if unknown

class EditMode(str, Enum):
    FAST = "fast"
    DETAILED = "detailed"

class EditRequest(BaseModel):
    text: str
    model: str
    mode: EditMode = EditMode.FAST

class EditResponse(BaseModel):
    edited_text: str
    technical_explanation: str

# Map frontend model names to backend model types
model_mapping = {
    "models/gemini-1.5-pro-latest": ModelType.GEMINI_PRO.value,
    "models/gemini-1.5-flash-8b": ModelType.GEMINI_FLASH.value,
    "gpt-3.5-turbo": ModelType.GPT35.value,
    "gpt-4": ModelType.GPT4.value,
    "gemini-pro": ModelType.GEMINI_PRO.value,  # For backward compatibility
    "gemini-flash": ModelType.GEMINI_FLASH.value  # For backward compatibility
}

def validate_word_count(original_text: str, edited_text: str, tolerance: int = 50) -> bool:
    """Validate that the edited text's word count is within tolerance of the original."""
    original_words = len(original_text.split())
    edited_words = len(edited_text.split())
    difference = abs(original_words - edited_words)
    return difference <= tolerance

async def process_gemini_edit(text: str, mode: EditMode = EditMode.FAST) -> str:
    """Process text editing using Gemini API with chunking for long texts."""
    try:
        logger.info("Starting Gemini edit process")
        model = get_gemini_model(ModelType.GEMINI_PRO.value)
        
        # Define the base prompt based on mode
        base_prompt = """شما یک ویراستار متخصص متون کوچینگ هستید. متن زیر را با حفظ معنا و کیفیت ویرایش کنید.

دستورالعمل‌های دقیق:
۱. متن را با دقت ویرایش کنید و تمام خطاها را اصلاح کنید
۲. تعداد کلمات متن ویرایش شده باید تقریباً برابر با متن اصلی باشد (حداکثر ۱۰۰ کلمه تفاوت)
۳. عناوین و سرفصل‌ها را دقیقاً حفظ کنید
۴. پاراگراف‌بندی را حفظ کنید
۵. اصلاحات ضروری را انجام دهید:
   - اصلاح خطاهای دستوری و املایی
   - بهبود ساختار جملات
   - حذف تکرار و حشو
   - بهینه‌سازی جملات
   - اصلاح خطاهای نگارشی
   - بهبود نقطه‌گذاری
   - اصلاح فاصله‌گذاری
   - بهبود روانی متن
   - اصلاح خطاهای منطقی
۶. لحن حرفه‌ای و رسمی را حفظ کنید
۷. اصطلاحات تخصصی کوچینگ را حفظ کنید
۸. معنا و محتوای اصلی را حفظ کنید
۹. از زبان رسمی و معیار استفاده کنید
۱۰. جملات را روان و واضح کنید
۱۱. حفظ پیوستگی و انسجام متن
۱۲. اصلاح خطاهای زمانی و مکانی در متن
۱۳. حفظ پیوستگی بین بخش‌های مختلف متن
۱۴. اطمینان از انسجام معنایی بین جملات

نکته مهم: هدف اصلی بهبود کیفیت متن و اصلاح خطاهاست. اگر خطایی در متن وجود دارد، حتماً آن را اصلاح کنید. در متون روایی و داستانی، حفظ روانی و انسجام متن از اهمیت بالایی برخوردار است. این متن بخشی از یک متن بزرگتر است، پس لطفاً پیوستگی معنایی و ساختاری آن را حفظ کنید."""
        
        # Split text into chunks if needed (max 400 words per chunk)
        words = text.split()
        chunks = []
        current_chunk = []
        current_word_count = 0
        max_words_per_chunk = 400
        
        for word in words:
            current_chunk.append(word)
            current_word_count += 1
            if current_word_count >= max_words_per_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_word_count = 0
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        logger.info(f"Split text into {len(chunks)} chunks")
        
        edited_chunks = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            chunk_prompt = f"{base_prompt}\n\nمتن برای ویرایش (بخش {i+1}/{len(chunks)}):\n{chunk}\n\nمتن ویرایش شده:"
            
            max_attempts = 3
            for attempt in range(max_attempts):
                try:
                    logger.info(f"Attempt {attempt + 1}/{max_attempts}")
                    response = model.generate_content(chunk_prompt)
                
                    if not response or not response.text:
                        logger.warning(f"Empty response on attempt {attempt + 1}")
                        continue
                    
                    edited_chunk = response.text.strip()
                    
                    # Log the differences between original and edited text
                    original_words = len(chunk.split())
                    edited_words = len(edited_chunk.split())
                    logger.info(f"Original words: {original_words}, Edited words: {edited_words}")
                    
                    # Validate word count for this chunk with more flexible tolerance
                    if validate_word_count(chunk, edited_chunk, tolerance=100):
                        logger.info(f"Word count validation passed for chunk {i+1}")
                        edited_chunks.append(edited_chunk)
                        break
                    else:
                        logger.warning(f"Word count validation failed for chunk {i+1} on attempt {attempt + 1}")
                        
                    if attempt == max_attempts - 1:
                        logger.warning(f"All attempts failed for chunk {i+1}, using original chunk")
                        edited_chunks.append(chunk)
                
                except Exception as e:
                    logger.error(f"Error on attempt {attempt + 1}: {str(e)}")
                    if attempt == max_attempts - 1:
                        raise Exception(f"Failed to process chunk {i+1} after {max_attempts} attempts: {str(e)}")
                    continue
                
        # Combine edited chunks with proper spacing and ensure continuity
        combined_text = '\n\n'.join(edited_chunks)
        
        # Clean up any double newlines and ensure proper spacing
        combined_text = re.sub(r'\n\s*\n', '\n\n', combined_text)
        
        return combined_text
            
    except Exception as e:
        logger.error(f"Gemini edit error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Gemini edit failed: {str(e)}")

async def chunk_text(text: str, max_chunk_size: int = 1500) -> List[str]:
    """Split text into chunks while preserving all text elements."""
    try:
        # Normalize line breaks while preserving intentional multiple breaks
        text = re.sub(r'([^\n])\n([^\n])', r'\1\n\n\2', text)  # Convert single newlines to double
        text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize multiple newlines to double
        
        # Split text into paragraphs preserving empty lines
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = []
        current_size = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            
            # Always preserve empty paragraphs
            if not paragraph:
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                chunks.append('')
                continue
            
            # If paragraph fits in current chunk
            if current_size + len(paragraph) + 2 <= max_chunk_size:
                current_chunk.append(paragraph)
                current_size += len(paragraph) + 2
            else:
                # Save current chunk if not empty
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Handle paragraphs larger than max_chunk_size
                if len(paragraph) > max_chunk_size:
                    # Split into sentences
                    sentence_pattern = re.compile(r'([.!?؟،;؛]+\s*)')
                    sentences = sentence_pattern.split(paragraph)
                    
                    temp_chunk = []
                    temp_size = 0
                    
                    for sentence in sentences:
                        if not sentence:
                            continue
                            
                        if temp_size + len(sentence) <= max_chunk_size:
                            temp_chunk.append(sentence)
                            temp_size += len(sentence)
                        else:
                            if temp_chunk:
                                chunks.append(''.join(temp_chunk))
                            temp_chunk = [sentence]
                            temp_size = len(sentence)
                    
                    if temp_chunk:
                        chunks.append(''.join(temp_chunk))
                else:
                    current_chunk = [paragraph]
                    current_size = len(paragraph)
        
        # Add final chunk
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks
        
    except Exception as e:
        logger.error(f"Error in chunk_text: {str(e)}")
        raise

async def process_openai_edit(text: str, model: str = ModelType.GPT35.value, mode: EditMode = EditMode.FAST) -> str:
    """Process text editing using OpenAI API with chunking for long texts."""
    try:
        logger.info(f"Starting OpenAI edit process with model: {model}")
        client = get_async_openai_client()
        
        # Increase chunk sizes for better handling of long texts
        max_chunk_size = 2500 if model == "gpt-4" else 2000  # Increased chunk sizes
        chunks = await chunk_text(text, max_chunk_size)
        logger.info(f"Split text into {len(chunks)} chunks")
        
        edited_chunks = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            try:
                # Different prompts for fast vs detailed editing
                if mode == EditMode.FAST:
                    system_prompt = """شما یک ویراستار متخصص متون کوچینگ در یک انتشارات حرفه‌ای هستید. هدف شما بهبود کیفیت متن با حفظ معنی، محتوای اصلی و اصطلاحات تخصصی کوچینگ است.
لطفاً متن را با رعایت موارد زیر ویرایش کنید:
۱. حفظ دقیق ساختار پاراگراف‌های متن اصلی
۲. حفظ تعداد پاراگراف‌ها و فاصله‌گذاری بین آنها
۳. اصلاح خطاهای دستوری و نقطه‌گذاری
۴. حفظ و استفاده صحیح از اصطلاحات تخصصی کوچینگ
۵. بهبود ساختار جملات و روانی متن
۶. حفظ دقیق معنا و محتوای اصلی
۷. رعایت اصول نگارش فارسی معیار
۸. تأکید بر حفظ و تقویت لحن رسمی و حرفه‌ای

نکته مهم: تعداد کلمات متن ویرایش شده باید تقریباً برابر با متن اصلی باشد (حداکثر ۲۰ کلمه کمتر یا بیشتر).
توجه: فقط ویرایش متن فارسی، بدون ترجمه."""
                else:
                    system_prompt = """شما یک ویراستار متخصص متون کوچینگ در یک انتشارات حرفه‌ای هستید. هدف شما بهبود کیفیت متن با حفظ معنی، محتوای اصلی و اصطلاحات تخصصی کوچینگ است.
لطفاً متن را با دقت و با رعایت موارد زیر ویرایش کنید:
۱. حفظ دقیق ساختار پاراگراف‌های متن اصلی
۲. حفظ تعداد پاراگراف‌ها و فاصله‌گذاری بین آنها
۳. اصلاح تمام خطاهای دستوری، املایی و نقطه‌گذاری
۴. حفظ و کاربرد دقیق اصطلاحات تخصصی کوچینگ
۵. بهبود ساختار جملات، وضوح و خوانایی متن
۶. ارتقای سطح نگارش و حرفه‌ای‌تر کردن متن
۷. حفظ دقیق معنا و محتوای اصلی
۸. رعایت اصول نگارش فارسی معیار و زبان رسمی
۹. تأکید ویژه بر حفظ و تقویت لحن رسمی و کاملاً حرفه‌ای

نکته مهم: تعداد کلمات متن ویرایش شده باید تقریباً برابر با متن اصلی باشد (حداکثر ۲۰ کلمه کمتر یا بیشتر).
توجه: فقط ویرایش متن فارسی، بدون ترجمه."""

                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"""لطفاً این متن کوچینگ (بخش {i+1}/{len(chunks)}) را با حفظ معنا، محتوا و اصطلاحات تخصصی آن ویرایش کنید.
تعداد کلمات متن ویرایش شده باید تقریباً برابر با متن اصلی باشد.

{chunk}

نکته مهم: فقط ویرایش متن فارسی، بدون ترجمه."""}
                ]
                
                max_attempts = 2  # Reduced from 3 to 2 attempts
                for attempt in range(max_attempts):
                    completion = await client.chat.completions.create(
                        model=model,
                        messages=messages,
                        temperature=0.1,  # Lower temperature for faster, more consistent responses
                        max_tokens=max_chunk_size,
                        timeout=20.0,  # Further reduced timeout
                        presence_penalty=-0.1,  # Slight penalty to prevent wordiness
                        frequency_penalty=0.1,  # Slight penalty to prevent repetition
                    )
                    
                    if not completion or not completion.choices:
                        logger.error(f"Empty response from OpenAI for chunk {i+1} on attempt {attempt + 1}")
                        continue
                    
                    edited_chunk = completion.choices[0].message.content.strip()
                    
                    # Validate word count for this chunk
                    if validate_word_count(chunk, edited_chunk):
                        edited_chunks.append(edited_chunk)
                        break
                    else:
                        logger.warning(f"Word count validation failed for chunk {i+1} on attempt {attempt + 1}")
                        
                    if attempt == max_attempts - 1:
                        logger.warning(f"All attempts failed for chunk {i+1}, using original chunk")
                        edited_chunks.append(chunk)
                
            except Exception as e:
                logger.error(f"Error processing chunk {i+1}: {str(e)}")
                edited_chunks.append(chunk)  # Keep original chunk if processing fails
                continue
        
        # Combine edited chunks with proper spacing
        edited_text = ''
        for i, chunk in enumerate(edited_chunks):
            chunk = chunk.strip()
            if not chunk:
                # Only add empty line if it was in the original text
                if i > 0 and i < len(edited_chunks) - 1 and edited_chunks[i-1].strip() and edited_chunks[i+1].strip():
                    edited_text += '\n\n'
                continue
            
            if i > 0 and edited_chunks[i-1].strip():
                # Add newline only between non-empty chunks
                edited_text += '\n\n'
            edited_text += chunk
        
        # Clean up any extra newlines while preserving paragraph structure
        edited_text = re.sub(r'\n\s*\n\s*\n', '\n\n', edited_text)
        edited_text = edited_text.strip()
        
        return edited_text
            
    except Exception as e:
        logger.error(f"OpenAI edit error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Edit failed: {str(e)}")

@app.post("/edit")
async def edit_text(request: EditRequest) -> EditResponse:
    """Edit Persian text for improved clarity and formality."""
    try:
        text = request.text.strip()
        model_type = model_mapping.get(request.model)
        mode = request.mode
        
        if not text:
            raise HTTPException(status_code=400, detail="Text cannot be empty")
            
        if not model_type:
            raise HTTPException(status_code=400, detail=f"Invalid model type: {request.model}")
            
        # Update length validation for 3000 words
        max_length = 25000 if model_type == ModelType.GPT4.value else 21000  # Approximately 3000-3500 words
        if len(text) > max_length:
            raise HTTPException(
                status_code=400, 
                detail=f"Text too long (max {max_length} characters, approximately 3000-3500 words for {model_type})"
            )
            
        try:
            if model_type in [ModelType.GEMINI_FLASH.value, ModelType.GEMINI_PRO.value]:
                edited = await process_gemini_edit(text, mode)
            else:
                # Ensure we're using the correct model type for OpenAI
                if model_type == ModelType.GPT35.value:
                    model_type = "gpt-3.5-turbo-0125"  # Use the latest version
                elif model_type == ModelType.GPT4.value:
                    model_type = "gpt-4-0125-preview"  # Use the latest version
                edited = await process_openai_edit(text, model_type, mode)
            
            # Validate edited text
            if not edited or not edited.strip():
                raise Exception("Received empty response from API")
            
            # Create diff and ensure proper encoding
            diff_html = create_diff_html(text, edited)
            
            # Ensure proper encoding of Persian text
            return JSONResponse(
                content={
                    "edited_text": edited.strip(),
                    "technical_explanation": f"✅ Used {request.model} for {mode.value} editing",
                    "diff_html": diff_html,
                    "model_used": model_type
                },
                headers={
                    "Content-Type": "application/json; charset=utf-8",
                    "Access-Control-Allow-Origin": "*"
                }
            )
            
        except Exception as e:
            logger.error(f"API processing error: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process text with {request.model}: {str(e)}"
            )
            
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Edit error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Unexpected error during text editing: {str(e)}"
        )

def create_diff_html(original: str, edited: str) -> str:
    """Create HTML with highlighted differences between original and edited text."""
    html = []
    changes = {
        'replacements': 0,
        'deletions': 0,
        'insertions': 0
    }
    
    html.append("""
    <style>
        .diff-container {
            direction: rtl;
            text-align: right;
            padding: 1em;
            font-family: inherit;
            line-height: 2;
            white-space: normal !important;
        }
        .text-content {
            display: block;
            white-space: normal !important;
            word-spacing: 0.25em;
        }
        .word {
            display: inline;
            white-space: normal !important;
            margin: 0 0.25em;
        }
        .delete {
            color: #ff0000 !important;
            background-color: #ffebee;
            text-decoration: line-through;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            display: inline;
            white-space: normal !important;
            margin: 0 0.25em;
        }
        .insert {
            color: #008000;
            background-color: #e8f5e9;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            display: inline;
            white-space: normal !important;
            margin: 0 0.25em;
        }
        .arrow {
            color: #666;
            margin: 0 0.5em;
            font-size: 0.9em;
            display: inline;
            white-space: normal !important;
        }
        .unchanged {
            display: inline;
            white-space: normal !important;
            margin: 0 0.25em;
        }
        .paragraph-break {
            display: block;
            margin: 1em 0;
            white-space: normal !important;
        }
        .space {
            display: inline;
            margin: 0 0.25em;
            white-space: normal !important;
        }
    </style>
    <div class="diff-container"><div class="text-content">""")
    
    # Split both texts into paragraphs while preserving empty lines
    def split_into_paragraphs(text):
        # Split by double newlines or more
        paragraphs = re.split(r'\n\s*\n', text)
        # Filter out empty paragraphs and strip whitespace
        return [p.strip() for p in paragraphs if p.strip()]
    
    original_paragraphs = split_into_paragraphs(original)
    edited_paragraphs = split_into_paragraphs(edited)
    
    # Ensure both lists have the same length by padding with empty strings
    max_length = max(len(original_paragraphs), len(edited_paragraphs))
    original_paragraphs.extend([''] * (max_length - len(original_paragraphs)))
    edited_paragraphs.extend([''] * (max_length - len(edited_paragraphs)))
    
    # Process each paragraph pair
    for i, (orig_para, edit_para) in enumerate(zip(original_paragraphs, edited_paragraphs)):
        html.append(f'<div class="paragraph-container">')
        
        # Split paragraphs into words
        orig_words = orig_para.split() if orig_para else []
        edit_words = edit_para.split() if edit_para else []
        
        # Create sequence matcher for word-level diff
        matcher = SequenceMatcher(None, orig_words, edit_words)
        
        # Process each operation in the diff
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                # Words that are the same in both versions
                for word in orig_words[i1:i2]:
                    html.append(f'<span class="unchanged">{escape(word)}</span>')
            elif op == 'delete':
                # Words that were deleted
                for word in orig_words[i1:i2]:
                    html.append(f'<span class="delete">{escape(word)}</span>')
                changes['deletions'] += 1
            elif op == 'insert':
                # Words that were added
                for word in edit_words[j1:j2]:
                    html.append(f'<span class="insert">{escape(word)}</span>')
                changes['insertions'] += 1
            elif op == 'replace':
                # Words that were replaced
                for word in orig_words[i1:i2]:
                    html.append(f'<span class="delete">{escape(word)}</span>')
                for word in edit_words[j1:j2]:
                    html.append(f'<span class="insert">{escape(word)}</span>')
                changes['replacements'] += 1
        
        html.append(f'</div>')
    
    # Calculate word counts
    orig_word_count = sum(len(p.split()) for p in original_paragraphs)
    edit_word_count = sum(len(p.split()) for p in edited_paragraphs)
    word_diff = edit_word_count - orig_word_count
    total_changes = changes['replacements'] + changes['deletions'] + changes['insertions']
    
    html.append("</div></div>")
    
    # Add changes summary
    html.append(f"""
    <div class="changes-summary">
        خلاصه تغییرات:
        <br>
        📊 تعداد کلمات متن اصلی: <span class="word-count">{orig_word_count}</span>
        <br>
        📊 تعداد کلمات متن ویرایش شده: <span class="word-count">{edit_word_count}</span>
        <br>
        📊 تفاوت تعداد کلمات: <span class="word-count">{word_diff:+d}</span>
        <br>
        <br>
        🔄 تعداد کل تغییرات: {total_changes}
        <br>
        🔀 جایگزینی‌ها: {changes['replacements']}
        <br>
        ❌ حذف‌ها: {changes['deletions']}
        <br>
        ✅ اضافه‌ها: {changes['insertions']}
    </div>
    """)
    
    return '\n'.join(html)

@app.post("/translate")
async def translate_text(request: Request):
    """Translate Persian text to English with support for long texts."""
    try:
        data = await request.json()
        logger.info(f"Received translation request with data: {data}")
        
        # Use edited_text if present, otherwise fall back to text
        text = data.get("edited_text", "").strip() or data.get("text", "").strip()
        model_type = data.get("model", "gpt-3.5-turbo")
        
        logger.info(f"Starting translation with model: {model_type}")
        logger.info(f"Text length: {len(text)} characters")
        
        if not text:
            logger.error("Empty text provided")
            raise HTTPException(status_code=400, detail="No text provided")
            
        # Validate model type
        valid_models = ["gpt-3.5-turbo", "gpt-4", "models/gemini-1.5-pro-latest", "models/gemini-1.5-flash-8b"]
        if model_type not in valid_models:
            logger.error(f"Invalid model type: {model_type}. Valid models are: {valid_models}")
            raise HTTPException(status_code=400, detail=f"Invalid model type: {model_type}")
        
        # Process translation based on model
        try:
            if model_type in ["gpt-3.5-turbo", "gpt-4"]:
                logger.info(f"Using OpenAI model: {model_type}")
                try:
                    client = get_async_openai_client()
                    
                    # Adjust chunk size based on model
                    if model_type == "gpt-4":
                        chunk_size = 4000  # Larger chunks for GPT-4
                    else:
                        chunk_size = 3000  # Increased chunk size for GPT-3.5
                    
                    # Split text into chunks while preserving structure
                    chunks = await chunk_text(text, chunk_size)
                    
                    # Log chunk information
                    logger.info(f"Split text into {len(chunks)} chunks")
                    for i, chunk in enumerate(chunks):
                        logger.info(f"Chunk {i+1} length: {len(chunk)} characters")
                        if chunk:
                            logger.info(f"Chunk {i+1} preview: {chunk[:50]}...")
                    
                    # Translate each chunk with overlap handling
                    translated_chunks = []
                    overlap_buffer = ""  # Store end of previous chunk
                    
                    for i, chunk in enumerate(chunks):
                        try:
                            logger.info(f"Translating chunk {i+1}/{len(chunks)}")
                            
                            # Skip empty chunks but preserve them
                            if not chunk.strip():
                                translated_chunks.append("")
                                continue
                            
                            # Add overlap buffer to beginning of chunk if exists
                            chunk_to_translate = overlap_buffer + chunk if overlap_buffer else chunk
                            
                            # Create system message with specific instructions
                            system_message = """You are a professional Persian to English translator.
                            Follow these rules strictly:
                            1. Preserve all formatting, paragraph breaks, and line breaks exactly
                            2. Maintain all punctuation marks in their correct positions
                            3. Ensure no words or phrases are missed
                            4. Keep the same paragraph structure
                            5. Only output the translation, nothing else
                            6. Preserve any special characters or symbols
                            7. Maintain consistent translation of repeated terms"""
                            
                            response = await client.chat.completions.create(
                                model=model_type,
                                messages=[
                                    {"role": "system", "content": system_message},
                                    {"role": "user", "content": f"Translate this Persian text to English:\n\n{chunk_to_translate}"}
                                ],
                                temperature=0.3,
                                presence_penalty=0.0,
                                frequency_penalty=0.0
                            )
                            
                            translated_chunk = response.choices[0].message.content.strip()
                            
                            # Store last sentence for overlap
                            if i < len(chunks) - 1:
                                last_sentence = re.split(r'([.!?])\s+', chunk)[-1]
                                overlap_buffer = last_sentence if len(last_sentence) < 100 else ""
                            
                            # Log translation result
                            logger.info(f"Successfully translated chunk {i+1}")
                            logger.info(f"Chunk {i+1} translation length: {len(translated_chunk)} characters")
                            if translated_chunk:
                                logger.info(f"Chunk {i+1} translation preview: {translated_chunk[:50]}...")
                            
                            translated_chunks.append(translated_chunk)
                            
                        except Exception as e:
                            logger.error(f"Error translating chunk {i+1}: {str(e)}")
                            raise HTTPException(status_code=500, detail=f"Translation error in chunk {i+1}: {str(e)}")
                    
                    # Combine translated chunks with proper spacing
                    translated_text = ""
                    for i, chunk in enumerate(translated_chunks):
                        chunk = chunk.strip()
                        if not chunk:
                            if i > 0 and i < len(translated_chunks) - 1:
                                translated_text += "\n\n"
                            continue
                        
                        if translated_text and not translated_text.endswith("\n"):
                            translated_text += "\n\n"
                        translated_text += chunk
                    
                    # Clean up the final text while preserving structure
                    translated_text = re.sub(r'\n{3,}', '\n\n', translated_text)
                    translated_text = translated_text.strip()
                    
                    # Verify translation completeness
                    original_words = len(text.split())
                    translated_words = len(translated_text.split())
                    logger.info(f"Original words: {original_words}, Translated words: {translated_words}")
                    
                    if abs(translated_words - original_words) > original_words * 0.4:  # Allow 40% difference due to language differences
                        logger.warning(f"Significant word count difference detected: {abs(translated_words - original_words)} words")
                    
                    return JSONResponse(content={
                        "translated_text": translated_text,
                        "model_used": model_type,
                        "original_length": len(text),
                        "translated_length": len(translated_text),
                        "original_word_count": original_words,
                        "translated_word_count": translated_words,
                        "chunk_count": len(chunks)
                    })
                    
                except Exception as e:
                    logger.error(f"OpenAI translation error: {str(e)}")
                    raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")
            
            elif model_type in ["models/gemini-1.5-pro-latest", "models/gemini-1.5-flash-8b"]:
                logger.info("Using Gemini model")
                try:
                    model = get_gemini_model(model_type)
                    
                    # Adjust chunk size based on model
                    if model_type == "gemini":
                        chunk_size = 1000  # Smaller chunks for Gemini
                    else:
                        chunk_size = 1500  # More conservative chunk size for GPT models
                    
                    # Split text into chunks while preserving paragraph structure
                    chunks = []
                    current_chunk = ""
                    paragraphs = text.split('\n\n')
                    
                    for paragraph in paragraphs:
                        if len(current_chunk) + len(paragraph) + 2 <= chunk_size:
                            current_chunk += paragraph + '\n\n'
                        else:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            current_chunk = paragraph + '\n\n'
                    
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    
                    # Log chunk information
                    logger.info(f"Split text into {len(chunks)} chunks")
                    for i, chunk in enumerate(chunks):
                        logger.info(f"Chunk {i+1} length: {len(chunk)} characters")
                    
                    # Translate each chunk
                    translated_chunks = []
                    for i, chunk in enumerate(chunks):
                        try:
                            logger.info(f"Translating chunk {i+1}/{len(chunks)} with Gemini")
                            if model_type == "gemini":
                                response = model.generate_content(
                                    f"""Translate the following Persian text to English. 
                                    Preserve all formatting, paragraph breaks, and line breaks exactly as they appear.
                                    Only output the translation, nothing else.

                                    Text to translate:
                                    {chunk}"""
                                )
                                translated_chunk = response.text.strip()
                            else:
                                response = await client.chat.completions.create(
                                    model="gpt-4",
                                    messages=[
                                        {"role": "system", "content": "You are a professional Persian to English translator. Preserve all formatting, paragraph breaks, and line breaks exactly as they appear in the original text. Only output the translation, nothing else."},
                                        {"role": "user", "content": f"Translate this Persian text to English:\n\n{chunk}"}
                                    ],
                                    temperature=0.3
                                )
                                translated_chunk = response.choices[0].message.content.strip()
                            
                            logger.info(f"Successfully translated chunk {i+1}")
                            translated_chunks.append(translated_chunk)
                        except Exception as e:
                            logger.error(f"Error translating chunk {i+1} with Gemini: {str(e)}")
                            raise HTTPException(status_code=500, detail=f"Error translating chunk {i+1}: {str(e)}")
                    
                    # Combine translated chunks
                    translated_text = ''
                    for chunk in translated_chunks:
                        chunk = chunk.strip()
                        if chunk:
                            # Remove any line breaks and extra spaces
                            chunk = chunk.replace('\n', ' ')
                            chunk = re.sub(r'\s+', ' ', chunk)
                            if translated_text:
                                translated_text += ' '
                            translated_text += chunk
                    
                    # Clean up the final text
                    translated_text = re.sub(r'\s+', ' ', translated_text)
                    translated_text = translated_text.strip()
                    
                    logger.info("Successfully completed Gemini translation")
                except Exception as e:
                    logger.error(f"Gemini API error: {str(e)}")
                    raise HTTPException(status_code=500, detail=f"Gemini API error: {str(e)}")
            
            else:
                logger.error(f"Invalid model type: {model_type}")
                raise HTTPException(status_code=400, detail=f"Invalid model type: {model_type}")

            if not translated_text:
                logger.error("Empty translation result")
                raise HTTPException(status_code=500, detail="Empty translation result")

            logger.info("Successfully completed translation")
            return JSONResponse(
                content={
                    "translated_text": translated_text,
                    "model_used": model_type
                },
                headers={
                    "Content-Type": "application/json; charset=utf-8"
                }
            )

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Translation API error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/export-word")
async def export_to_word(request: Request):
    try:
        data = await request.json()
        original_text = data.get("original_text", "")
        edited_text = data.get("edited_text", "")
        combined_view_html = data.get("combined_view", "")
        
        # Create a new Word document
        doc = Document()
        
        # Set page margins to be smaller
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
        
        # Add title
        title = doc.add_heading('Persian Text Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add original text section
        doc.add_heading('Original Text', level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(original_text)
        run.font.rtl = True
        run.font.size = Pt(12)
        
        # Add edited text section
        doc.add_heading('Edited Text', level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(edited_text)
        run.font.rtl = True
        run.font.size = Pt(12)
        
        # Add combined view section
        doc.add_heading('Combined View with Changes', level=1)
        current_paragraph = doc.add_paragraph()
        current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Parse the HTML content
        soup = BeautifulSoup(combined_view_html, 'html.parser')
        
        def add_text_with_format(paragraph, text, class_name=None):
            """Add text to paragraph with proper formatting."""
            if not text:
                return

            # Split text into words while preserving spaces
            words = text.split()
            
            # Process each word
            for word in words:
                # Add the word
                run = paragraph.add_run(word)
                run.font.rtl = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Always black text
                
                # Apply special formatting based on class
                if class_name == 'delete':
                    run.font.strike = True
                    # Use custom shading for light red
                    run._r.get_or_add_rPr().append(
                        parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFEBEE"/>')
                    )
                elif class_name == 'insert':
                    # Use custom shading for light green (matching paragraph-by-paragraph view)
                    run._r.get_or_add_rPr().append(
                        parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="E8F5E9"/>')
                    )
                elif class_name == 'arrow':
                    pass  # Keep arrows in black
                
                # Add space after word
                space_run = paragraph.add_run(' ')
                space_run.font.rtl = True
                space_run.font.size = Pt(12)
                space_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Apply the same highlighting to the space
                if class_name == 'delete':
                    space_run.font.strike = True
                    # Use custom shading for light red
                    space_run._r.get_or_add_rPr().append(
                        parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFEBEE"/>')
                    )
                elif class_name == 'insert':
                    # Use custom shading for light green (matching paragraph-by-paragraph view)
                    space_run._r.get_or_add_rPr().append(
                        parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="E8F5E9"/>')
                    )
        
        # Find the main content container
        content_div = soup.find('div', class_='diff-container')
        if content_div:
            text_content = content_div.find('div', class_='text-content')
            if text_content:
                current_paragraph = doc.add_paragraph()
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Process each element
                for element in text_content.find_all(['span', 'div']):
                    # Get element's class
                    class_name = element.get('class', [None])[0]
                    
                    if class_name in ['delete', 'insert', 'unchanged', 'arrow']:
                        # Get text content
                        text = element.get_text()
                        if text:
                            # Add the text with formatting
                            add_text_with_format(current_paragraph, text, class_name)
                    elif class_name == 'paragraph-break':
                        # Create new paragraph for line breaks
                        current_paragraph = doc.add_paragraph()
                        current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add paragraph-by-paragraph comparison section
        doc.add_heading('Paragraph-by-Paragraph Comparison', level=1)
        
        def splitIntoChunks(text):
            """Split text into chunks of maximum 400 characters, trying to break at sentence endings."""
            if not text:
                return []
                
            chunks = []
            current_chunk = ""
            sentences = re.split('([.!?])', text)
            
            for i in range(0, len(sentences), 2):
                sentence = sentences[i] + (sentences[i + 1] if i + 1 < len(sentences) else '')
                if len(current_chunk) + len(sentence) <= 400:
                    current_chunk += sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence
                    
            if current_chunk:
                chunks.append(current_chunk.strip())
                
            return chunks
        
        # Split texts into chunks using the same logic as UI
        original_chunks = splitIntoChunks(original_text)
        edited_chunks = splitIntoChunks(edited_text)
        
        # Create comparison for each chunk
        for i in range(max(len(original_chunks), len(edited_chunks))):
            # Add chunk number
            doc.add_heading(f'Chunk {i + 1}', level=2)
            
            # Original chunk
            doc.add_heading('Original', level=3)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(original_chunks[i] if i < len(original_chunks) else '')
            run.font.rtl = True
            run.font.size = Pt(12)
            
            # Edited chunk with light green background
            doc.add_heading('Edited', level=3)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p._p.get_or_add_pPr().append(
                parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="E8F5E9"/>')
            )
            run = p.add_run(edited_chunks[i] if i < len(edited_chunks) else original_chunks[i])
            run.font.rtl = True
            run.font.size = Pt(12)
            
            # Add spacing between chunks
            doc.add_paragraph()
        
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Return the document
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=persian_text.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def read_root(request: Request):
    """Serve the main page."""
    try:
        return templates.TemplateResponse("index.html", {
            "request": request
        })
    except Exception as e:
        logger.error(f"Error serving index page: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/healthz")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import signal
    import sys
    import asyncio
    
    def signal_handler(sig, frame):
        logger.info("Received shutdown signal, cleaning up...")
        sys.exit(0)
    
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)
    
    try:
        # Initialize API clients
        if not asyncio.run(initialize_api_clients()):
            logger.error("Failed to initialize API clients. Exiting...")
            sys.exit(1)
        
        # Try ports in a higher range
        start_port = 9090
        end_port = 9100
        
        def is_port_in_use(port):
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                return s.connect_ex(('localhost', port)) == 0
                
        def find_available_port(start_port=9090, max_port=9100):
            for port in range(start_port, max_port + 1):
                if not is_port_in_use(port):
                    return port
            raise RuntimeError("No available ports found")
        
        port = find_available_port()
        
        try:
            logger.info(f"Starting server on port {port}")
            uvicorn.run(
                "translation_bot:app",  # Pass as import string instead of app object
                host="0.0.0.0",
                port=port,
                reload=True,
                log_level="info",
                access_log=True,
                workers=1
            )
        except Exception as e:
            logger.error(f"Failed to start server on port {port}: {str(e)}")
            sys.exit(1)
                
    except Exception as e:
        logger.error(f"Server startup failed: {str(e)}")
        sys.exit(1) 